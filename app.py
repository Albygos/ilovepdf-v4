import os
import uuid
import time
import zipfile
import shutil
import io
import threading
import re
from flask import Flask, request, jsonify, send_file, send_from_directory, redirect, abort
from pypdf import PdfReader, PdfWriter
import fitz  # PyMuPDF
from PIL import Image
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

from pptx import Presentation
from pptx.util import Inches, Pt
import openpyxl
from reportlab.platypus import Table as RLTable, TableStyle as RLTableStyle, PageBreak
from reportlab.lib import colors
from deep_translator import GoogleTranslator

from pdf_tools import *



app = Flask(__name__)

# Setup directories inside workspace
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
PROCESSED_FOLDER = os.path.join(BASE_DIR, 'processed')

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROCESSED_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['PROCESSED_FOLDER'] = PROCESSED_FOLDER
# Limit maximum upload to 50MB
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

def clean_old_files():
    """Deletes files in uploads and processed folders older than 15 minutes."""
    while True:
        now = time.time()
        for folder in [UPLOAD_FOLDER, PROCESSED_FOLDER]:
            if os.path.exists(folder):
                for filename in os.listdir(folder):
                    file_path = os.path.join(folder, filename)
                    try:
                        # Skip special system files
                        if filename.startswith('.'):
                            continue
                        if os.path.isfile(file_path) or os.path.islink(file_path):
                            if os.path.getmtime(file_path) < now - 900:  # 15 minutes
                                os.unlink(file_path)
                        elif os.path.isdir(file_path):
                            if os.path.getmtime(file_path) < now - 900:
                                shutil.rmtree(file_path)
                    except Exception as e:
                        print(f"Error cleaning file {file_path}: {e}")
        time.sleep(300)  # Check every 5 minutes

# Start the clean up thread as a daemon
cleanup_thread = threading.Thread(target=clean_old_files, daemon=True)
cleanup_thread.start()

from werkzeug.utils import secure_filename
import traceback

def get_unique_filename(ext):
    return f"{uuid.uuid4().hex}.{ext}"

from werkzeug.exceptions import HTTPException

@app.errorhandler(Exception)
def handle_global_exception(e):
    if isinstance(e, HTTPException):
        return e
    traceback.print_exc()
    if isinstance(e, ValueError):
        return jsonify({'error': str(e)}), 400
    return jsonify({'error': 'An internal server error occurred.'}), 500

def process_upload(request_files, expected_extensions, max_files=1):
    """
    Generic utility to validate and save uploaded files.
    expected_extensions should be a list like ['.pdf'] or ['.jpg', '.png']
    """
    if not request_files:
        raise ValueError('No files uploaded.')
        
    saved_paths = []
    
    # In some forms it's request.files.getlist('files'), in others (request.files.get('file') or request.files.get('files'))
    # This standardizes to a list.
    if isinstance(request_files, list):
        files = request_files
    else:
        files = [request_files]
        
    if len(files) > max_files:
        raise ValueError(f'Too many files uploaded. Maximum is {max_files}.')
        
    for f in files:
        if f.filename == '':
            continue
            
        ext = os.path.splitext(f.filename)[1].lower()
        if ext not in expected_extensions:
            raise ValueError(f'Invalid file type. Expected one of {", ".join(expected_extensions)}')
            
        secure_name = secure_filename(f.filename)
        # We append a UUID to prevent collisions but keep the safe original name (if any)
        final_name = f"{uuid.uuid4().hex}_{secure_name}" if secure_name else get_unique_filename(ext[1:])
        
        path = os.path.join(UPLOAD_FOLDER, final_name)
        f.save(path)
        saved_paths.append(path)
        
    if not saved_paths:
        raise ValueError('No valid files were processed.')
        
    return saved_paths



# Real Original ISO Language & Locale Validator (EXPANDED REGIONAL & COUNTRY CODES)
VALID_LANG_CODES_SET = set(['af', 'af-za', 'am', 'am-et', 'ar', 'ar-ae', 'ar-bh', 'ar-dz', 'ar-eg', 'ar-iq', 'ar-jo', 'ar-kw', 'ar-lb', 'ar-ly', 'ar-ma', 'ar-om', 'ar-qa', 'ar-sa', 'ar-sd', 'ar-sy', 'ar-tn', 'ar-ye', 'az', 'bg', 'bg-bg', 'bn', 'bn-bd', 'bn-in', 'bs', 'ca', 'ca-es', 'ceb', 'ceb-ph', 'co', 'cs', 'cs-cz', 'cy', 'cy-gb', 'da', 'da-dk', 'de', 'de-at', 'de-ch', 'de-de', 'de-li', 'de-lu', 'el', 'el-cy', 'el-gr', 'en', 'en-au', 'en-ca', 'en-gb', 'en-ie', 'en-in', 'en-nz', 'en-ph', 'en-sg', 'en-us', 'en-za', 'eo', 'es', 'es-ar', 'es-bo', 'es-cl', 'es-co', 'es-cr', 'es-cub', 'es-do', 'es-ec', 'es-es', 'es-gt', 'es-hn', 'es-mx', 'es-ni', 'es-pa', 'es-pe', 'es-pr', 'es-py', 'es-sv', 'es-uy', 'es-ve', 'et', 'et-ee', 'eu', 'eu-es', 'fa', 'fa-ir', 'fi', 'fi-fi', 'fil', 'fil-ph', 'fr', 'fr-be', 'fr-ca', 'fr-ch', 'fr-ci', 'fr-cm', 'fr-dz', 'fr-fr', 'fr-lu', 'fr-mc', 'fr-sn', 'fy', 'ga', 'ga-ie', 'gd', 'gl', 'gl-es', 'gu', 'gu-in', 'ha', 'ha-ng', 'haw', 'haw-us', 'he', 'hi', 'hi-in', 'hr', 'hr-hr', 'ht', 'ht-ht', 'hu', 'hu-hu', 'hy', 'id', 'id-id', 'ig', 'ig-ng', 'is', 'is-is', 'it', 'it-ch', 'it-it', 'ja', 'ja-jp', 'jv', 'jv-id', 'ka', 'kk', 'km', 'km-kh', 'kn', 'kn-in', 'ko', 'ko-kr', 'la', 'lb', 'lo', 'lo-la', 'lt', 'lt-lt', 'lv', 'lv-lv', 'mg', 'mg-mg', 'mi', 'mi-nz', 'mk', 'mk-mk', 'ml', 'ml-in', 'mn', 'mr', 'mr-in', 'ms', 'ms-my', 'ms-sg', 'mt', 'mt-mt', 'my', 'my-mm', 'nb-no', 'ne', 'ne-np', 'nl', 'nl-be', 'nl-nl', 'nn-no', 'no-no', 'ny', 'ny-mw', 'pa', 'pa-in', 'pa-pk', 'pl', 'pl-pl', 'ps', 'ps-af', 'pt', 'pt-ao', 'pt-br', 'pt-mz', 'pt-pt', 'ro', 'ro-md', 'ro-ro', 'ru', 'ru-by', 'ru-kz', 'ru-ru', 'sd', 'sd-pk', 'si', 'si-lk', 'sk', 'sk-sk', 'sl', 'sl-si', 'sm', 'sm-ws', 'so', 'so-so', 'sq', 'sq-al', 'sr', 'sr-rs', 'st', 'st-ls', 'su', 'su-id', 'sv', 'sv-fi', 'sv-se', 'sw', 'sw-ke', 'sw-tz', 'ta', 'ta-in', 'ta-lk', 'ta-sg', 'te', 'te-in', 'th', 'th-th', 'tn', 'tn-bw', 'tr', 'tr-tr', 'uk', 'uk-ua', 'ur', 'ur-in', 'ur-pk', 'uz', 've', 'vi', 'vi-vn', 'xh', 'xh-za', 'yo', 'yo-ng', 'zh', 'zh-cn', 'zh-hk', 'zh-mo', 'zh-sg', 'zh-tw', 'zu', 'zu-za'])

def is_valid_language_code(code):
    if not code:
        return False
    code_lower = code.lower()
    return code_lower in VALID_LANG_CODES_SET

SUPPORTED_LANGS = {'en': 'English', 'es': 'Español', 'fr': 'Français', 'de': 'Deutsch', 'pt': 'Português', 'hi': 'हिन्दी', 'ar': 'العربية', 'zh': '中文', 'ja': '日本語', 'ko': '한국어', 'ru': 'Русский', 'it': 'Italiano', 'nl': 'Nederlands', 'tr': 'Türkçe', 'pl': 'Polski', 'sv': 'Svenska', 'id': 'Bahasa Indonesia', 'vi': 'Tiếng Việt', 'th': 'ไทย', 'uk': 'Українська', 'el': 'Ελληνικά', 'he': 'עברית', 'bn': 'বাংলা', 'ta': 'தமிழ்', 'te': 'తెలుగు', 'mr': 'मराठी', 'gu': 'ગુજરાતી', 'kn': 'કન્નડ', 'ur': 'اردو', 'sw': 'Kiswahili', 'ms': 'Bahasa Melayu', 'fil': 'Filipino', 'cs': 'Čeština', 'hu': 'Magyar', 'da': 'Dansk', 'fi': 'Suomi', 'sk': 'Slovenčina', 'ro': 'Română', 'bg': 'Български', 'hr': 'Hrvatski', 'sr': 'Српски', 'sl': 'Slovenščina', 'et': 'Eesti', 'lv': 'Latviešu', 'lt': 'Lietuvių', 'sq': 'Shqip', 'mk': 'Македонски', 'ka': 'ქართული', 'hy': 'Հայերեն', 'az': 'Azərbaycan', 'kk': 'Қазақша', 'uz': 'O‘zbekcha', 'mn': 'Монгол', 'my': 'မြန်မာ', 'km': 'ភាសាខ្មែរ', 'lo': 'ພາສາລາວ', 'si': 'සිංහල', 'ne': 'नेपाली', 'pa': 'ਪੰਜਾਬੀ', 'am': 'አማርኛ', 'so': 'Soomaali', 'yo': 'Yorùbá', 'ig': 'Asụsụ Igbo', 'ha': 'Hausa', 'zu': 'isiZulu', 'xh': 'isiXhosa', 'af': 'Afrikaans', 'gl': 'Galego', 'ca': 'Català', 'eu': 'Euskara', 'cy': 'Cymraeg', 'ga': 'Gaeilge', 'is': 'Íslenska', 'mt': 'Malti', 'lb': 'Lëtzebuergesch', 'bs': 'Bosanski', 'fa': 'فارسی', 'ps': 'پښتو', 'sd': 'سنڌي', 'su': 'Basa Sunda', 'jv': 'Basa Jawa', 'ceb': 'Cebuano', 'haw': 'ʻŌlelo Hawaiʻi', 'eo': 'Esperanto', 'la': 'Latina', 'gd': 'Gàidhlig', 'fy': 'Frysk', 'co': 'Corsu', 'mi': 'Te Reo Māori', 'sm': 'Gagana Sāmoa', 'ny': 'Chichewa', 'ht': 'Kreyòl Ayisyen', 'mg': 'Malagasy', 'st': 'Sesotho', 'tn': 'Setswana', 've': 'Tshivenḓa'}

SLUG_TO_FILE = {
    '': 'index.html',
    'merge_pdf': 'merge.html',
    'split_pdf': 'split.html',
    'compress_pdf': 'compress.html',
    'word_to_pdf': 'word-to-pdf.html',
    'pdf_to_word': 'pdf-to-word.html',
    'jpg_to_pdf': 'jpg-to-pdf.html',
    'pdf_to_jpg': 'pdf-to-jpg.html',
    'rotate_pdf': 'rotate.html',
    'protect-pdf': 'protect.html',
    'unlock_pdf': 'unlock.html',
    'pdf_add_watermark': 'watermark.html',
    'add_pdf_page_number': 'page-numbers.html',
    'organize-pdf': 'organize.html',
    'html-to-pdf': 'html-to-pdf.html',
    'extract-text': 'extract-text.html',
    'remove-pages': 'delete-pages.html',
    'compare-pdf': 'compare.html',
    'powerpoint_to_pdf': 'powerpoint-to-pdf.html',
    'powerpoint-to-pdf': 'powerpoint-to-pdf.html',
    'pdf_to_powerpoint': 'pdf-to-powerpoint.html',
    'pdf-to-powerpoint': 'pdf-to-powerpoint.html',
    'excel_to_pdf': 'excel-to-pdf.html',
    'excel-to-pdf': 'excel-to-pdf.html',
    'pdf_to_excel': 'pdf-to-excel.html',
    'pdf-to-excel': 'pdf-to-excel.html',
    'convert-pdf-to-pdfa': 'pdf-to-pdfa.html',
    'repair-pdf': 'repair-pdf.html',
    'ocr-pdf': 'ocr-pdf.html',
    'pdf-summarize': 'pdf-summarize.html',
    'translate-pdf': 'translate-pdf.html',
    'sign-pdf': 'sign-pdf.html',
    'privacy-policy': 'privacy-policy.html',
    'terms': 'terms.html',
    'about': 'about.html',
    'contact': 'contact.html',
    'cookie-policy': 'cookie-policy.html',
    'disclaimer': 'disclaimer.html'
}

LEGACY_REDIRECTS = {
    'merge': 'merge_pdf',
    'split': 'split_pdf',
    'compress': 'compress_pdf',
    'word-to-pdf': 'word_to_pdf',
    'pdf-to-word': 'pdf_to_word',
    'jpg-to-pdf': 'jpg_to_pdf',
    'pdf-to-jpg': 'pdf_to_jpg',
    'rotate': 'rotate_pdf',
    'protect': 'protect-pdf',
    'unlock': 'unlock_pdf',
    'watermark': 'pdf_add_watermark',
    'page-numbers': 'add_pdf_page_number',
    'organize': 'organize-pdf',
    'delete-pages': 'remove-pages',
    'compare': 'compare-pdf',
    'html_to_pdf': 'html-to-pdf',
    'extract_text': 'extract-text',
    'remove_pages': 'remove-pages',
    'compare_pdf': 'compare-pdf',
    'powerpoint_to_pdf': 'powerpoint-to-pdf',
    'pdf_to_powerpoint': 'pdf-to-powerpoint',
    'excel_to_pdf': 'excel-to-pdf',
    'pdf_to_excel': 'pdf-to-excel',
    'repair_pdf': 'repair-pdf',
    'ocr_pdf': 'ocr-pdf',
    'translate_pdf': 'translate-pdf',
    'sign_pdf': 'sign-pdf'
}

TRANSLATIONS = {'es': {'Every tool you need to work with PDFs in one place': 'Todas las herramientas PDF que necesitas en un solo lugar', 'Every tool you need to use PDFs, at your fingertips. All are 100% FREE and easy to use! Merge, split, compress, convert, rotate, unlock and watermark PDFs with just a few clicks.': 'Todas las herramientas PDF que necesitas al alcance de tu mano. ¡100% GRATIS y fáciles de usar! Une, divide, comprime, convierte, rota, desbloquea y añade marcas de agua a tus archivos PDF con solo unos pocos clics.', 'Merge PDF': 'Unir PDF', 'Split PDF': 'Dividir PDF', 'Compress PDF': 'Comprimir PDF', 'Word to PDF': 'Word a PDF', 'PDF to Word': 'PDF a Word', 'JPG to PDF': 'JPG a PDF', 'PDF to JPG': 'PDF a JPG', 'Rotate PDF': 'Rotar PDF', 'Protect PDF': 'Proteger PDF', 'Unlock PDF': 'Desbloquear PDF', 'Watermark': 'Marca de agua', 'Page Numbers': 'Números de página', 'Organize PDF': 'Organizar PDF', 'HTML to PDF': 'HTML a PDF', 'Extract Text': 'Extraer texto', 'Delete Pages': 'Eliminar páginas', 'PDF to PDF/A': 'PDF a PDF/A', 'Repair PDF': 'Reparar PDF', 'OCR PDF': 'OCR PDF', 'AI Summarizer': 'Resumidor de IA', 'Translate PDF': 'Traducir PDF', 'Sign PDF': 'Firmar PDF', 'Compare PDF': 'Comparar PDF', 'All PDF Tools': 'Todas las herramientas PDF', 'Convert PDF': 'Convertir PDF', 'Log in': 'Iniciar sesión', 'Sign up': 'Registrarse', 'Home': 'Inicio'}, 'fr': {'Every tool you need to work with PDFs in one place': 'Tous les outils PDF dont vous avez besoin au même endroit', 'Every tool you need to use PDFs, at your fingertips. All are 100% FREE and easy to use! Merge, split, compress, convert, rotate, unlock and watermark PDFs with just a few clicks.': 'Tous les outils PDF à portée de main. 100% GRATUITS et faciles à utiliser ! Fusionnez, divisez, compressez, convertissez, pivotez, déverrouillez et ajoutez un filigrane à vos PDFs en quelques clics.', 'Merge PDF': 'Fusionner PDF', 'Split PDF': 'Diviser PDF', 'Compress PDF': 'Compresser PDF', 'Word to PDF': 'Word en PDF', 'PDF to Word': 'PDF en Word', 'JPG to PDF': 'JPG en PDF', 'PDF to JPG': 'PDF en JPG', 'Rotate PDF': 'Pivoter PDF', 'Protect PDF': 'Protéger PDF', 'Unlock PDF': 'Déverrouiller PDF', 'Watermark': 'Filigrane', 'Page Numbers': 'Numérotation', 'Organize PDF': 'Organiser PDF', 'HTML to PDF': 'HTML en PDF', 'Extract Text': 'Extraire le texte', 'Delete Pages': 'Supprimer des pages', 'PDF to PDF/A': 'PDF en PDF/A', 'Repair PDF': 'Réparer PDF', 'OCR PDF': 'OCR PDF', 'AI Summarizer': 'Résumeur IA', 'Translate PDF': 'Traduire PDF', 'Sign PDF': 'Signer PDF', 'Compare PDF': 'Comparer PDF', 'All PDF Tools': 'Tous les outils PDF', 'Convert PDF': 'Convertir PDF', 'Log in': 'Connexion', 'Sign up': "S'inscrire", 'Home': 'Accueil'}, 'de': {'Every tool you need to work with PDFs in one place': 'Jedes PDF-Werkzeug, das Sie benötigen, an einem Ort', 'Every tool you need to use PDFs, at your fingertips. All are 100% FREE and easy to use! Merge, split, compress, convert, rotate, unlock and watermark PDFs with just a few clicks.': 'Alle PDF-Werkzeuge für Sie bereit. 100% KOSTENLOS und einfach zu bedienen! PDF-Dateien zusammenfügen, teilen, komprimieren, konvertieren, drehen, entsperren und Wasserzeichen hinzufügen.', 'Merge PDF': 'PDF zusammenfügen', 'Split PDF': 'PDF teilen', 'Compress PDF': 'PDF komprimieren', 'Word to PDF': 'Word in PDF', 'PDF to Word': 'PDF in Word', 'JPG to PDF': 'JPG in PDF', 'PDF to JPG': 'PDF in JPG', 'Rotate PDF': 'PDF drehen', 'Protect PDF': 'PDF schützen', 'Unlock PDF': 'PDF entsperren', 'Watermark': 'Wasserzeichen', 'Page Numbers': 'Seitenzahlen', 'Organize PDF': 'PDF organisieren', 'HTML to PDF': 'HTML in PDF', 'Extract Text': 'Text extrahieren', 'Delete Pages': 'Seiten löschen', 'PDF to PDF/A': 'PDF in PDF/A', 'Repair PDF': 'PDF reparieren', 'OCR PDF': 'OCR PDF', 'AI Summarizer': 'KI-Zusammenfassung', 'Translate PDF': 'PDF übersetzen', 'Sign PDF': 'PDF unterschreiben', 'Compare PDF': 'PDF vergleichen', 'All PDF Tools': 'Alle PDF-Werkzeuge', 'Convert PDF': 'PDF konvertieren', 'Log in': 'Einloggen', 'Sign up': 'Registrieren', 'Home': 'Startseite'}, 'pt': {'Every tool you need to work with PDFs in one place': 'Todas as ferramentas PDF necessárias em um só lugar', 'Every tool you need to use PDFs, at your fingertips. All are 100% FREE and easy to use! Merge, split, compress, convert, rotate, unlock and watermark PDFs with just a few clicks.': "Todas as ferramentas PDF ao seu alcance. 100% GRATUITAS e fáceis de usar! Mesclar, dividir, comprimir, converter, rotacionar, desbloquear e adicionar marcas d'água aos seus PDFs.", 'Merge PDF': 'Mesclar PDF', 'Split PDF': 'Dividir PDF', 'Compress PDF': 'Comprimir PDF', 'Word to PDF': 'Word para PDF', 'PDF to Word': 'PDF para Word', 'JPG to PDF': 'JPG para PDF', 'PDF to JPG': 'PDF para JPG', 'Rotate PDF': 'Rotacionar PDF', 'Protect PDF': 'Proteger PDF', 'Unlock PDF': 'Desbloquear PDF', 'Watermark': "Marca d'água", 'Page Numbers': 'Números de página', 'Organize PDF': 'Organizar PDF', 'HTML to PDF': 'HTML para PDF', 'Extract Text': 'Extrair texto', 'Delete Pages': 'Excluir páginas', 'PDF to PDF/A': 'PDF para PDF/A', 'Repair PDF': 'Reparar PDF', 'OCR PDF': 'OCR PDF', 'AI Summarizer': 'Resumidor de IA', 'Translate PDF': 'Traduzir PDF', 'Sign PDF': 'Assinar PDF', 'Compare PDF': 'Comparar PDF', 'All PDF Tools': 'Todas as ferramentas PDF', 'Convert PDF': 'Converter PDF', 'Log in': 'Entrar', 'Sign up': 'Cadastrar', 'Home': 'Início'}, 'hi': {'Every tool you need to work with PDFs in one place': 'पीडीएफ पर काम करने के लिए सभी आवश्यक उपकरण एक जगह', 'Every tool you need to use PDFs, at your fingertips. All are 100% FREE and easy to use! Merge, split, compress, convert, rotate, unlock and watermark PDFs with just a few clicks.': 'आपकी उंगलियों पर सभी पीडीएफ उपकरण। 100% मुफ़्त और उपयोग में आसान! कुछ ही क्लिक में पीडीएफ़ को मर्ज, स्प्लिट, कंप्रेस, कन्वर्ट, रोटेट, अनलॉक और वॉटरमार्क करें।', 'Merge PDF': 'पीडीएफ मर्ज करें', 'Split PDF': 'पीडीएफ विभाजित करें', 'Compress PDF': 'पीडीएफ कंप्रेस करें', 'Word to PDF': 'वर्ड से पीडीएफ', 'PDF to Word': 'पीडीएफ से वर्ड', 'JPG to PDF': 'जेपीजी से पीडीएफ', 'PDF to JPG': 'पीडीएफ से जेपीजी', 'Rotate PDF': 'पीडीएफ घुमाएँ', 'Protect PDF': 'पीडीएफ सुरक्षित करें', 'Unlock PDF': 'पीडीएफ अनलॉक करें', 'Watermark': 'वॉटरमार्क', 'Page Numbers': 'पेज नंबर', 'Organize PDF': 'पीडीएफ व्यवस्थित करें', 'HTML to PDF': 'एचटीएमएल से पीडीएफ', 'Extract Text': 'टेक्स्ट निकालें', 'Delete Pages': 'पेज हटाएं', 'PDF to PDF/A': 'पीडीएफ से पीडीएफ/ए', 'Repair PDF': 'पीडीएफ रिपेयर करें', 'OCR PDF': 'ओसीआर पीडीएफ', 'AI Summarizer': 'एआई सारांश', 'Translate PDF': 'अनुवाद करें', 'Sign PDF': 'हस्ताक्षर करें', 'Compare PDF': 'पीडीएफ तुलना करें', 'All PDF Tools': 'सभी पीडीएफ टूल्स', 'Convert PDF': 'कन्वर्ट पीडीएफ', 'Log in': 'लॉगिन', 'Sign up': 'साइन अप', 'Home': 'होम'}, 'ar': {'Every tool you need to work with PDFs in one place': 'جميع أدوات PDF التي تحتاجها في مكان واحد', 'Merge PDF': 'دمج PDF', 'Split PDF': 'تقسيم PDF', 'Compress PDF': 'ضغط PDF', 'Word to PDF': 'Word إلى PDF', 'PDF to Word': 'PDF إلى Word', 'JPG to PDF': 'JPG إلى PDF', 'PDF to JPG': 'PDF إلى JPG', 'Rotate PDF': 'تدوير PDF', 'Protect PDF': 'حماية PDF', 'Unlock PDF': 'فتح PDF', 'Watermark': 'علامة مائية', 'Page Numbers': 'أرقام الصفحات', 'Organize PDF': 'تنظيم PDF', 'HTML to PDF': 'HTML إلى PDF', 'Extract Text': 'استخراج النص', 'Delete Pages': 'حذف الصفحات', 'PDF to PDF/A': 'PDF إلى PDF/A', 'Repair PDF': 'إصلاح PDF', 'OCR PDF': 'التعرف الضوئي', 'AI Summarizer': 'ملخص الذكاء الاصطناعي', 'Translate PDF': 'ترجمة PDF', 'Sign PDF': 'توقيع PDF', 'Compare PDF': 'مقارنة PDF', 'All PDF Tools': 'جميع أدوات PDF', 'Convert PDF': 'تحويل PDF', 'Log in': 'تسجيل الدخول', 'Sign up': 'إنشاء حساب', 'Home': 'الرئيسية'}, 'zh': {'Every tool you need to work with PDFs in one place': '您在一个地方所需的所有 PDF 工具', 'Merge PDF': '合并 PDF', 'Split PDF': '拆分 PDF', 'Compress PDF': '压缩 PDF', 'Word to PDF': 'Word 转 PDF', 'PDF to Word': 'PDF 转 Word', 'JPG to PDF': 'JPG 转 PDF', 'PDF to JPG': 'PDF 转 JPG', 'Rotate PDF': '旋转 PDF', 'Protect PDF': '加密 PDF', 'Unlock PDF': '解密 PDF', 'Watermark': '添加水印', 'Page Numbers': '添加页码', 'Organize PDF': '组织 PDF', 'HTML to PDF': 'HTML 转 PDF', 'Extract Text': '提取文本', 'Delete Pages': '删除页面', 'PDF to PDF/A': 'PDF 转 PDF/A', 'Repair PDF': '修复 PDF', 'OCR PDF': 'OCR 识别', 'AI Summarizer': 'AI 摘要', 'Translate PDF': '翻译 PDF', 'Sign PDF': '电子签名', 'Compare PDF': '对比 PDF', 'All PDF Tools': '所有 PDF 工具', 'Convert PDF': '转换 PDF', 'Log in': '登录', 'Sign up': '注册', 'Home': '首页'}, 'ja': {'Every tool you need to work with PDFs in one place': 'PDFの作業に必要なすべてのツールを1箇所に', 'Merge PDF': 'PDF 結合', 'Split PDF': 'PDF 分割', 'Compress PDF': 'PDF 圧縮', 'Word to PDF': 'Word から PDF', 'PDF to Word': 'PDF から Word', 'JPG to PDF': 'JPG から PDF', 'PDF to JPG': 'PDF から JPG', 'Rotate PDF': 'PDF 回転', 'Protect PDF': 'PDF 保護', 'Unlock PDF': 'PDF ロック解除', 'Watermark': '透かし追加', 'Page Numbers': 'ページ番号', 'Organize PDF': 'PDF 整理', 'HTML to PDF': 'HTML から PDF', 'Extract Text': 'テキスト抽出', 'Delete Pages': 'ページ削除', 'PDF to PDF/A': 'PDF から PDF/A', 'Repair PDF': 'PDF 修復', 'OCR PDF': 'OCR 処理', 'AI Summarizer': 'AI 要約', 'Translate PDF': 'PDF 翻訳', 'Sign PDF': 'PDF 署名', 'Compare PDF': 'PDF 比較', 'All PDF Tools': 'すべての PDF ツール', 'Convert PDF': 'PDF 変換', 'Log in': 'ログイン', 'Sign up': '新規登録', 'Home': 'ホーム'}, 'ru': {'Every tool you need to work with PDFs in one place': 'Все инструменты PDF, необходимые для работы, в одном месте', 'Merge PDF': 'Объединить PDF', 'Split PDF': 'Разделить PDF', 'Compress PDF': 'Сжать PDF', 'Word to PDF': 'Word в PDF', 'PDF to Word': 'PDF в Word', 'JPG to PDF': 'JPG в PDF', 'PDF to JPG': 'PDF в JPG', 'Rotate PDF': 'Повернуть PDF', 'Protect PDF': 'Защитить PDF', 'Unlock PDF': 'Снять защиту PDF', 'Watermark': 'Водяной знак', 'Page Numbers': 'Номера страниц', 'Organize PDF': 'Упорядочить PDF', 'HTML to PDF': 'HTML в PDF', 'Extract Text': 'Извлечь текст', 'Delete Pages': 'Удалить страницы', 'PDF to PDF/A': 'PDF в PDF/A', 'Repair PDF': 'Восстановить PDF', 'OCR PDF': 'Распознавание OCR', 'AI Summarizer': 'ИИ Суммаризатор', 'Translate PDF': 'Перевести PDF', 'Sign PDF': 'Подписать PDF', 'Compare PDF': 'Сравнить PDF', 'All PDF Tools': 'Все инструменты PDF', 'Convert PDF': 'Конвертировать PDF', 'Log in': 'Войти', 'Sign up': 'Регистрация', 'Home': 'Главная'}}
# Serve explicit static assets from root
@app.route('/style.css')
def serve_css():
    return send_from_directory(BASE_DIR, 'style.css')

@app.route('/ad-manager.js')
def serve_ad_manager():
    return send_from_directory(BASE_DIR, 'ad-manager.js')

@app.route('/site.webmanifest')
def serve_webmanifest():
    return send_from_directory(BASE_DIR, 'site.webmanifest')

@app.route('/favicon.ico')
def serve_favicon_ico():
    return send_from_directory(BASE_DIR, 'favicon.ico')

@app.route('/favicon-32x32.png')
def serve_favicon_32():
    return send_from_directory(BASE_DIR, 'favicon-32x32.png')

@app.route('/favicon-16x16.png')
def serve_favicon_16():
    return send_from_directory(BASE_DIR, 'favicon-16x16.png')

@app.route('/favicon-192x192.png')
def serve_favicon_192():
    return send_from_directory(BASE_DIR, 'favicon-192x192.png')

@app.route('/apple-touch-icon.png')
def serve_apple_icon():
    return send_from_directory(BASE_DIR, 'apple-touch-icon.png')

@app.route('/logo-branding.png')
def serve_logo_branding():
    return send_from_directory(BASE_DIR, 'logo-branding.png')

@app.route('/ads.txt')
def serve_ads_txt():
    return send_from_directory(BASE_DIR, 'ads.txt')

@app.route('/privacy-policy')
def serve_privacy_policy():
    return send_from_directory(BASE_DIR, 'privacy-policy.html')

@app.route('/terms')
def serve_terms():
    return send_from_directory(BASE_DIR, 'terms.html')

@app.route('/about')
def serve_about():
    return send_from_directory(BASE_DIR, 'about.html')

@app.route('/contact')
def serve_contact():
    return send_from_directory(BASE_DIR, 'contact.html')

@app.route('/cookie-policy')
def serve_cookie_policy():
    return send_from_directory(BASE_DIR, 'cookie-policy.html')

@app.route('/disclaimer')
def serve_disclaimer():
    return send_from_directory(BASE_DIR, 'disclaimer.html')

# Serve pages cleanly matching URLs and localized parameters
@app.route('/')
@app.route('/<lang>')
def index_route(lang=None):
    if lang and is_valid_language_code(lang):
        return serve_tool_page('', lang)
    if lang and (lang in SLUG_TO_FILE or lang in LEGACY_REDIRECTS):
        return serve_tool_page(lang, None)
    if lang:
        return dynamic_seo_page(lang)
    return serve_tool_page('', None)

@app.route('/<lang>/<slug>')
def lang_tool_route(lang, slug):
    if is_valid_language_code(lang) and (slug in SLUG_TO_FILE or slug in LEGACY_REDIRECTS):
        return serve_tool_page(slug, lang)
    elif lang in SLUG_TO_FILE or lang in LEGACY_REDIRECTS:
        return serve_tool_page(lang, None)
    return dynamic_seo_page(f"{lang}/{slug}")

def serve_tool_page(slug, lang=None):
    # Backward compatibility redirects
    if slug in LEGACY_REDIRECTS:
        try:
            target = LEGACY_REDIRECTS[slug]
            url = f"/{lang}/{target}" if lang else f"/{target}"
            return redirect(url, code=301)
        except Exception as e:
            import traceback
            return traceback.format_exc(), 500

    filename = SLUG_TO_FILE.get(slug)
    if not filename:
        # Fallback check if first parameter was actually slug without lang
        if slug in SUPPORTED_LANGS:
            filename = 'index.html'
        else:
            return "Page not found.", 404
            
    filepath = os.path.join(BASE_DIR, filename)
    if not os.path.exists(filepath):
        return "File not found.", 404
        
    with open(filepath, 'r', encoding='utf-8') as f:
        html = f.read()
        
    # --- Inject Canonical & Hreflang SEO Metadata ---
    root_domain = "https://ilovespdfs.in"
    lang_sub = f"{lang}/" if lang and lang != 'en' else ""
    canonical_url = f"{root_domain}/{lang_sub}{slug}"
    
    hreflangs = []
    hreflangs.append(f'<link rel="alternate" hreflang="x-default" href="{root_domain}/{slug}" />')
    hreflangs.append(f'<link rel="alternate" hreflang="en" href="{root_domain}/{slug}" />')
    indian_langs_set = {'hi', 'bn', 'ta', 'te', 'mr', 'gu', 'kn', 'ur', 'pa', 'ml', 'or'}
    for l in SUPPORTED_LANGS:
        if l != 'en':
            hreflangs.append(f'<link rel="alternate" hreflang="{l}" href="{root_domain}/{l}/{slug}" />')
            if l in indian_langs_set:
                hreflangs.append(f'<link rel="alternate" hreflang="{l}-in" href="{root_domain}/{l}/{slug}" />')
            
    hreflang_tags = "\n    ".join(hreflangs)
    
    # Strip existing canonical and inject new
    import re
    html = re.sub(r'<link\s+rel=["\']canonical["\'].*?>', '', html, flags=re.IGNORECASE)
    injection = f'<link rel="canonical" href="{canonical_url}" />\n    {hreflang_tags}'
    html = re.sub(r'(<head\b[^>]*>)', r'\1\n    ' + injection, html, flags=re.IGNORECASE)
    
    # --- Server-Side Dictionary Translation ---
    if lang and lang in TRANSLATIONS and lang != 'en':
        trans_dict = TRANSLATIONS[lang]
        sorted_keys = sorted(trans_dict.keys(), key=len, reverse=True)
        for key in sorted_keys:
            html = html.replace(key, trans_dict[key])
            
    return apply_language_translations(html, lang)

# Compatibility fallbacks
@app.route('/merge')
@app.route('/split')
@app.route('/compress')
@app.route('/word-to-pdf')
@app.route('/pdf-to-word')
@app.route('/jpg-to-pdf')
@app.route('/pdf-to-jpg')
@app.route('/rotate')
@app.route('/protect')
@app.route('/unlock')
@app.route('/watermark')
@app.route('/page-numbers')
@app.route('/organize')
@app.route('/html-to-pdf')
@app.route('/extract-text')
@app.route('/delete-pages')
@app.route('/compare')
def legacy_page_redirect():
    slug = request.path.strip('/')
    return serve_tool_page(slug, None)

@app.errorhandler(413)
def request_entity_too_large(error):
    return jsonify({'error': 'File is too large. Maximum size is 50MB.'}), 413

# --- PDF Tools API ---

@app.route('/api/merge', methods=['POST'])
def merge_pdfs():
    try:
        saved_paths = process_upload(request.files.getlist('files'), ['.pdf'], max_files=100)
        if len(saved_paths) < 2:
            raise ValueError('At least 2 PDF files are required for merging.')
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    try:

        # Merge
        merger = PdfWriter()
        for path in saved_paths:
            merger.append(path)

        out_filename = f"merged_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        with open(out_path, 'wb') as f_out:
            merger.write(f_out)
        merger.close()

        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        # Clean up input uploads immediately
        for path in saved_paths:
            try:
                os.remove(path)
            except:
                pass

@app.route('/api/split', methods=['POST'])
def split_pdf():
    file = (request.files.get('file') or request.files.get('files'))
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400

    ranges_str = request.form.get('ranges', '1-end') # e.g. "1-3, 4-5" or "1-end"
    
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)

    try:
        reader = PdfReader(input_path)
        total_pages = len(reader.pages)

        # Parse ranges
        ranges = []
        parts = [p.strip() for p in ranges_str.split(',') if p.strip()]
        
        try:
            for part in parts:
                if '-' in part:
                    start_s, end_s = part.split('-', 1)
                    start_s = start_s.strip()
                    end_s = end_s.strip()
                    
                    start = int(start_s) if start_s else 1
                    if end_s.lower() == 'end' or not end_s:
                        end = total_pages
                    else:
                        end = int(end_s)
                else:
                    start = int(part)
                    end = start
                
                # Bound checking (1-based index)
                start = max(1, min(start, total_pages))
                end = max(1, min(end, total_pages))
                if start <= end:
                    ranges.append((start, end))
                else:
                    ranges.append((end, start))
        except ValueError:
            return jsonify({'error': 'Invalid range format. Please use integers like "1-3" or "1-end".'}), 400

        if not ranges:
            return jsonify({'error': 'Invalid ranges specified.'}), 400

        # If only one range, return single PDF
        if len(ranges) == 1:
            start, end = ranges[0]
            writer = PdfWriter()
            for p_num in range(start - 1, end):
                writer.add_page(reader.pages[p_num])
            
            out_filename = f"split_{start}-{end}_{uuid.uuid4().hex[:8]}.pdf"
            out_path = os.path.join(PROCESSED_FOLDER, out_filename)
            with open(out_path, 'wb') as f_out:
                writer.write(f_out)
            
            return jsonify({'download_url': f'/download/{out_filename}'})
        else:
            # If multiple ranges, create a ZIP
            zip_filename = f"split_{uuid.uuid4().hex[:8]}.zip"
            zip_path = os.path.join(PROCESSED_FOLDER, zip_filename)
            
            with zipfile.ZipFile(zip_path, 'w') as zf:
                for idx, (start, end) in enumerate(ranges):
                    writer = PdfWriter()
                    for p_num in range(start - 1, end):
                        writer.add_page(reader.pages[p_num])
                    
                    pdf_bytes = io.BytesIO()
                    writer.write(pdf_bytes)
                    pdf_bytes.seek(0)
                    zf.writestr(f"split_range_{idx+1}_{start}-{end}.pdf", pdf_bytes.read())

            return jsonify({'download_url': f'/download/{zip_filename}'})

    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/compress', methods=['POST'])
def compress_pdf():
    try:
        saved_paths = process_upload((request.files.get('file') or request.files.get('files')), ['.pdf'])
        input_path = saved_paths[0]
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    try:
        # Try PyMuPDF optimization first (highly effective)
        try:
            doc = fitz.open(input_path)
            out_filename = f"compressed_{uuid.uuid4().hex[:8]}.pdf"
            out_path = os.path.join(PROCESSED_FOLDER, out_filename)
            doc.save(out_path, garbage=4, deflate=True, clean=True)
            doc.close()
        except Exception:
            # Fallback to pypdf stream compression
            reader = PdfReader(input_path)
            writer = PdfWriter()
            for page in reader.pages:
                added_page = writer.add_page(page)
                added_page.compress_content_streams()
            out_filename = f"compressed_{uuid.uuid4().hex[:8]}.pdf"
            out_path = os.path.join(PROCESSED_FOLDER, out_filename)
            with open(out_path, 'wb') as f_out:
                writer.write(f_out)

        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/jpg_to_pdf', methods=['POST'])
def jpg_to_pdf():
    try:
        saved_paths = process_upload(request.files.getlist('files'), ['.jpg', '.jpeg', '.png', '.webp', '.bmp'], max_files=100)
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    try:

        # Convert images to PDF
        images = []
        for path in saved_paths:
            img = Image.open(path)
            if img.mode != 'RGB':
                img = img.convert('RGB')
            images.append(img)

        out_filename = f"converted_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        
        if images:
            images[0].save(out_path, save_all=True, append_images=images[1:])

        # Close all image handlers
        for img in images:
            img.close()

        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        for path in saved_paths:
            try:
                os.remove(path)
            except:
                pass

@app.route('/api/pdf_to_jpg', methods=['POST'])
def pdf_to_jpg():
    try:
        saved_paths = process_upload((request.files.get('file') or request.files.get('files')), ['.pdf'])
        input_path = saved_paths[0]
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    try:
        doc = fitz.open(input_path)
        zip_filename = f"images_{uuid.uuid4().hex[:8]}.zip"
        zip_path = os.path.join(PROCESSED_FOLDER, zip_filename)

        with zipfile.ZipFile(zip_path, 'w') as zf:
            for idx, page in enumerate(doc):
                # Render page to image (pixmap)
                pix = page.get_pixmap(dpi=150)
                # Convert colorspace to RGB if it is not RGB
                if pix.colorspace and pix.colorspace.n != 3:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                img_data = pix.tobytes("jpg")
                zf.writestr(f"page_{idx+1}.jpg", img_data)

        doc.close()
        return jsonify({'download_url': f'/download/{zip_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/rotate', methods=['POST'])
def rotate_pdf():
    file = (request.files.get('file') or request.files.get('files'))
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400

    try:
        angle = int(request.form.get('angle', 90))
        if angle not in [90, 180, 270, 360, -90, -180, -270]:
            angle = 90
    except ValueError:
        angle = 90

    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)

    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()

        for page in reader.pages:
            page.rotate(angle)
            writer.add_page(page)

        out_filename = f"rotated_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        with open(out_path, 'wb') as f_out:
            writer.write(f_out)

        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/protect', methods=['POST'])
def protect_pdf():
    file = (request.files.get('file') or request.files.get('files'))
    password = request.form.get('password')
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    if not password:
        return jsonify({'error': 'Password is required to protect the PDF.'}), 400

    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)

    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()
        writer.append_pages_from_reader(reader)
        writer.encrypt(password)

        out_filename = f"protected_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        with open(out_path, 'wb') as f_out:
            writer.write(f_out)

        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/unlock', methods=['POST'])
def unlock_pdf():
    file = (request.files.get('file') or request.files.get('files'))
    password = request.form.get('password', '')
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400

    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)

    try:
        reader = PdfReader(input_path)
        
        if reader.is_encrypted:
            success = reader.decrypt(password)
            if not success:
                return jsonify({'error': 'Incorrect password. Could not decrypt PDF.'}), 400
                
        writer = PdfWriter()
        writer.append_pages_from_reader(reader)

        out_filename = f"unlocked_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        with open(out_path, 'wb') as f_out:
            writer.write(f_out)

        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/watermark', methods=['POST'])
def watermark_pdf():
    file = (request.files.get('file') or request.files.get('files'))
    text = request.form.get('text', 'i loves pdf')
    try:
        opacity = float(request.form.get('opacity', 0.3))
        opacity = max(0.0, min(opacity, 1.0))
    except ValueError:
        opacity = 0.3
        
    try:
        rotation = int(request.form.get('rotation', 45))
    except ValueError:
        rotation = 45

    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400

    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)

    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()

        for page in reader.pages:
            box = page.cropbox
            width = float(box.width)
            height = float(box.height)
            
            # Generate watermark overlay PDF in memory matching page dimensions
            packet = io.BytesIO()
            can = canvas.Canvas(packet, pagesize=(width, height))
            can.setFont("Helvetica", 40)
            can.setFillColorRGB(0.7, 0.7, 0.7, alpha=opacity)
            can.saveState()
            
            # Center in page dynamically
            can.translate(width / 2, height / 2)
            can.rotate(rotation)
            can.drawCentredString(0, 0, text)
            can.restoreState()
            can.save()
            packet.seek(0)
            
            watermark_reader = PdfReader(packet)
            watermark_page = watermark_reader.pages[0]
            
            page.merge_page(watermark_page)
            writer.add_page(page)

        out_filename = f"watermarked_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        with open(out_path, 'wb') as f_out:
            writer.write(f_out)

        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/page_numbers', methods=['POST'])
def add_page_numbers():
    file = (request.files.get('file') or request.files.get('files'))
    position = request.form.get('position', 'bottom-right') # bottom-left, bottom-center, bottom-right, top-left, top-center, top-right
    try:
        start_num = int(request.form.get('start_number', 1))
    except ValueError:
        start_num = 1
    
    label_format = request.form.get('format', 'Page {n}') # e.g. "Page {n}" or "{n}"

    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400

    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)

    try:
        reader = PdfReader(input_path)
        total_pages = len(reader.pages)
        writer = PdfWriter()

        for idx in range(total_pages):
            page = reader.pages[idx]
            
            # Check cropbox size of the page
            box = page.cropbox
            width = float(box.width)
            height = float(box.height)
            
            # Generate overlay page number PDF
            packet = io.BytesIO()
            can = canvas.Canvas(packet, pagesize=(width, height))
            can.setFont("Helvetica", 10)
            can.setFillColorRGB(0.3, 0.3, 0.3)
            
            num_text = label_format.replace('{n}', str(start_num + idx))
            
            # Determine x, y based on margin and size
            margin = 36
            if 'top' in position:
                y = height - margin
            else:
                y = margin
                
            if 'left' in position:
                x = margin
                can.drawString(x, y, num_text)
            elif 'center' in position:
                x = width / 2
                can.drawCentredString(x, y, num_text)
            else: # right
                x = width - margin
                can.drawRightString(x, y, num_text)
                
            can.save()
            packet.seek(0)
            
            number_pdf = PdfReader(packet)
            page.merge_page(number_pdf.pages[0])
            writer.add_page(page)

        out_filename = f"numbered_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        with open(out_path, 'wb') as f_out:
            writer.write(f_out)

        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/organize', methods=['POST'])
def organize_pdf():
    file = (request.files.get('file') or request.files.get('files'))
    order_str = request.form.get('order')  # comma-separated 0-based index list e.g. "0,2,1"
    rotations_str = request.form.get('rotations') # comma-separated angles e.g. "90,0,180"
    
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    if not order_str:
        return jsonify({'error': 'Page ordering config is required.'}), 400

    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)

    try:
        reader = PdfReader(input_path)
        total_pages = len(reader.pages)
        
        try:
            indices = [int(i.strip()) for i in order_str.split(',') if i.strip()]
        except ValueError:
            return jsonify({'error': 'Invalid page ordering format.'}), 400

        rotations = []
        if rotations_str:
            try:
                rotations = [int(r.strip()) for r in rotations_str.split(',') if r.strip()]
            except ValueError:
                pass

        writer = PdfWriter()
        for idx_in_list, idx in enumerate(indices):
            if 0 <= idx < total_pages:
                page = writer.add_page(reader.pages[idx])
                if rotations and idx_in_list < len(rotations):
                    angle = rotations[idx_in_list]
                    if angle in [90, 180, 270, 360, -90, -180, -270]:
                        page.rotate(angle)

        out_filename = f"organized_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        with open(out_path, 'wb') as f_out:
            writer.write(f_out)

        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/word_to_pdf', methods=['POST'])
def word_to_pdf():
    try:
        saved_paths = process_upload((request.files.get('file') or request.files.get('files')), ['.doc', '.docx'])
        input_path = saved_paths[0]
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    try:
        # Load the Word document
        doc = Document(input_path)
        
        # Build ReportLab elements
        out_filename = f"converted_word_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        
        pdf_doc = SimpleDocTemplate(out_path, pagesize=letter,
                                    rightMargin=72, leftMargin=72,
                                    topMargin=72, bottomMargin=72)
        
        styles = getSampleStyleSheet()
        
        # Create a clean normal style
        normal_style = ParagraphStyle(
            name='WordNormal',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=11,
            leading=14,
            spaceAfter=8
        )
        
        heading1_style = ParagraphStyle(
            name='WordHeading1',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=18,
            leading=22,
            spaceAfter=12,
            spaceBefore=12
        )
        
        story = []
        
        from reportlab.platypus import Table as RLTable, TableStyle as RLTableStyle
        from reportlab.lib import colors
        
        for item in iter_block_items(doc):
            if isinstance(item, DocxParagraph):
                text = item.text.strip()
                if not text:
                    story.append(Spacer(1, 10))
                    continue
                
                style = normal_style
                if item.style and item.style.name:
                    if item.style.name.startswith('Heading 1'):
                        style = heading1_style
                    elif item.style.name.startswith('Heading'):
                        style = ParagraphStyle(
                            name='WordHeadingSub',
                            parent=styles['Heading2'],
                            fontName='Helvetica-Bold',
                            fontSize=14,
                            leading=18,
                            spaceAfter=10,
                            spaceBefore=10
                        )
                
                safe_text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe_text, style))
            elif isinstance(item, DocxTable):
                table_data = []
                for row in item.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = "\n".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                        safe_cell_text = cell_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        row_data.append(Paragraph(safe_cell_text or ' ', normal_style))
                    table_data.append(row_data)
                
                if table_data:
                    rl_table = RLTable(table_data, colWidths=None)
                    rl_table.setStyle(RLTableStyle([
                        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
                        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                        ('VALIGN', (0,0), (-1,-1), 'TOP'),
                        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
                        ('TOPPADDING', (0,0), (-1,-1), 6),
                        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
                        ('LEFTPADDING', (0,0), (-1,-1), 6),
                        ('RIGHTPADDING', (0,0), (-1,-1), 6),
                    ]))
                    story.append(rl_table)
                    story.append(Spacer(1, 12))
            
        # Build PDF file
        pdf_doc.build(story)

        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

# --- New PDF Tools API ---

@app.route('/api/pdf_to_word', methods=['POST'])
def pdf_to_word():
    try:
        saved_paths = process_upload((request.files.get('file') or request.files.get('files')), ['.pdf'])
        input_path = saved_paths[0]
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    try:
        doc = fitz.open(input_path)
        word_doc = Document()
        
        for page in doc:
            blocks = page.get_text("blocks")
            blocks.sort(key=lambda b: (b[1], b[0]))
            for block in blocks:
                block_text = block[4].strip()
                if block_text:
                    paragraph_text = " ".join([line.strip() for line in block_text.split("\n") if line.strip()])
                    word_doc.add_paragraph(paragraph_text)
        
        out_filename = f"converted_pdf_{uuid.uuid4().hex[:8]}.docx"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        word_doc.save(out_path)
        doc.close()
        
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/html_to_pdf', methods=['POST'])
def html_to_pdf():
    file = (request.files.get('file') or request.files.get('files'))
    html_code = request.form.get('html_code', '')

    html_content = ""
    if file and file.filename != '':
        # Read uploaded file
        try:
            html_content = file.read().decode('utf-8')
        except Exception as e:
            return jsonify({'error': f'Failed to read HTML file: {str(e)}'}), 400
    else:
        html_content = html_code

    if not html_content.strip():
        return jsonify({'error': 'HTML content or file is required.'}), 400

    out_filename = f"html_pdf_{uuid.uuid4().hex[:8]}.pdf"
    out_path = os.path.join(PROCESSED_FOLDER, out_filename)

    try:
        # Attempt conversion via fitz HTML support
        doc = fitz.open("html", html_content)
        pdf_bytes = doc.convert_to_pdf()
        with open(out_path, 'wb') as f_out:
            f_out.write(pdf_bytes)
        doc.close()
    except Exception as fitz_err:
        # Fallback to reportlab custom parser
        try:
            pdf_doc = SimpleDocTemplate(out_path, pagesize=letter)
            styles = getSampleStyleSheet()
            normal_style = ParagraphStyle(
                name='HTMLNormal',
                parent=styles['Normal'],
                fontName='Helvetica',
                fontSize=11,
                leading=14,
                spaceAfter=8
            )
            parser = ReportLabHTMLParser(styles, normal_style)
            parser.feed(html_content)
            parser.flush()
            if not parser.story:
                parser.story.append(Paragraph("Empty Document", normal_style))
            pdf_doc.build(parser.story)
        except Exception as rl_err:
            return jsonify({'error': f'PDF conversion failed: {str(fitz_err)} | Fallback: {str(rl_err)}'}), 500

    return jsonify({'download_url': f'/download/{out_filename}'})

@app.route('/api/extract_text', methods=['POST'])
def extract_text():
    try:
        saved_paths = process_upload((request.files.get('file') or request.files.get('files')), ['.pdf'])
        input_path = saved_paths[0]
    except ValueError as e:
        return jsonify({'error': str(e)}), 400

    try:
        doc = fitz.open(input_path)
        text_content = ""
        for page in doc:
            text_content += page.get_text() + "\n"
        
        out_filename = f"extracted_{uuid.uuid4().hex[:8]}.txt"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        with open(out_path, 'w', encoding='utf-8') as f_out:
            f_out.write(text_content)
        doc.close()
        
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/delete_pages', methods=['POST'])
def delete_pdf_pages():
    file = (request.files.get('file') or request.files.get('files'))
    pages_str = request.form.get('pages', '')
    
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    if not pages_str:
        return jsonify({'error': 'Pages to delete configuration is required.'}), 400

    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)

    try:
        reader = PdfReader(input_path)
        total_pages = len(reader.pages)
        
        # Parse ranges to delete
        deleted_indices = set()
        parts = [p.strip() for p in pages_str.split(',') if p.strip()]
        for part in parts:
            if '-' in part:
                start_s, end_s = part.split('-', 1)
                start = int(start_s) if start_s.strip() else 1
                if end_s.lower() == 'end' or not end_s.strip():
                    end = total_pages
                else:
                    end = int(end_s)
                for p in range(min(start, end), max(start, end) + 1):
                    if 1 <= p <= total_pages:
                        deleted_indices.add(p - 1)
            else:
                p = int(part)
                if 1 <= p <= total_pages:
                    deleted_indices.add(p - 1)
        
        if len(deleted_indices) >= total_pages:
            return jsonify({'error': 'Cannot delete all pages from the PDF.'}), 400
            
        writer = PdfWriter()
        for idx in range(total_pages):
            if idx not in deleted_indices:
                writer.add_page(reader.pages[idx])
                
        out_filename = f"deleted_pages_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        with open(out_path, 'wb') as f_out:
            writer.write(f_out)
            
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

# --- File Serving and Downloads ---

@app.route('/download/<filename>')
def download_file(filename):
    filename = os.path.basename(filename)
    file_path = os.path.join(PROCESSED_FOLDER, filename)
    if os.path.exists(file_path):
        return send_from_directory(PROCESSED_FOLDER, filename, as_attachment=True)
    else:
        return "File not found or expired.", 404

@app.route('/compare-pdf')
def compare_pdf_redirect():
    return serve_tool_page('compare-pdf', None)

# --- Technical SEO and Crawl Control ---

@app.route('/robots.txt')
def serve_robots():
    return send_from_directory(BASE_DIR, 'robots.txt', mimetype='text/plain')

@app.route('/sitemap.xml')
def serve_sitemap():
    return send_from_directory(BASE_DIR, 'sitemap.xml', mimetype='application/xml')


@app.after_request
def add_header(response):
    path = request.path
    if (path.endswith('.css') or path.endswith('.js') or 
        path.endswith('.png') or path.endswith('.jpg') or 
        path.endswith('.svg') or path.endswith('.webp') or 
        path.endswith('.ico') or path.endswith('.woff2')):
        response.headers['Cache-Control'] = 'public, max-age=31536000, immutable'
    else:
        response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    return response

@app.route('/api/powerpoint_to_pdf', methods=['POST'])
def powerpoint_to_pdf():
    file = (request.files.get('file') or request.files.get('files'))
    if not file or file.filename == '':
        return jsonify({'error': 'PowerPoint file is required.'}), 400
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pptx'))
    file.save(input_path)
    try:
        out_filename = f"converted_pptx_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        convert_pptx_to_pdf(input_path, out_path)
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/pdf_to_powerpoint', methods=['POST'])
def pdf_to_powerpoint():
    file = (request.files.get('file') or request.files.get('files'))
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)
    try:
        out_filename = f"converted_pdf_{uuid.uuid4().hex[:8]}.pptx"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        convert_pdf_to_pptx(input_path, out_path)
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/excel_to_pdf', methods=['POST'])
def excel_to_pdf():
    file = (request.files.get('file') or request.files.get('files'))
    if not file or file.filename == '':
        return jsonify({'error': 'Excel file is required.'}), 400
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('xlsx'))
    file.save(input_path)
    try:
        out_filename = f"converted_xlsx_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        convert_excel_to_pdf(input_path, out_path)
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/pdf_to_excel', methods=['POST'])
def pdf_to_excel():
    file = (request.files.get('file') or request.files.get('files'))
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)
    try:
        out_filename = f"converted_pdf_{uuid.uuid4().hex[:8]}.xlsx"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        convert_pdf_to_excel(input_path, out_path)
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/pdf_to_pdfa', methods=['POST'])
def pdf_to_pdfa():
    file = (request.files.get('file') or request.files.get('files'))
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)
    try:
        out_filename = f"archived_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        try:
            doc = fitz.open(input_path)
            doc.save(out_path, pdfa=True)
            doc.close()
        except Exception:
            shutil.copy(input_path, out_path)
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/repair_pdf', methods=['POST'])
def repair_pdf():
    file = (request.files.get('file') or request.files.get('files'))
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)
    try:
        out_filename = f"repaired_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        doc = fitz.open(input_path)
        doc.save(out_path, garbage=4, deflate=True, clean=True)
        doc.close()
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/ocr_pdf', methods=['POST'])
def ocr_pdf():
    file = (request.files.get('file') or request.files.get('files'))
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)
    try:
        out_filename = f"ocr_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        run_ocr_on_pdf(input_path, out_path)
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/pdf_summarize', methods=['POST'])
def pdf_summarize():
    file = (request.files.get('file') or request.files.get('files'))
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)
    try:
        out_filename = f"summary_{uuid.uuid4().hex[:8]}.txt"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        summarize_pdf_document(input_path, out_path)
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/translate_pdf', methods=['POST'])
def translate_pdf():
    file = (request.files.get('file') or request.files.get('files'))
    target_lang = request.form.get('target_lang', 'es')
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)
    try:
        out_filename = f"translated_{target_lang}_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        translate_pdf_document(input_path, out_path, target_lang)
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/api/sign_pdf', methods=['POST'])
def sign_pdf():
    file = (request.files.get('file') or request.files.get('files'))
    sig_b64 = request.form.get('signature_base64')
    try:
        page_idx = int(request.form.get('page_index', 0))
        left = float(request.form.get('left', 0.0))
        top = float(request.form.get('top', 0.0))
        width = float(request.form.get('width', 150.0))
        height = float(request.form.get('height', 60.0))
    except ValueError:
        return jsonify({'error': 'Invalid placement coordinates.'}), 400
    if not file or file.filename == '':
        return jsonify({'error': 'PDF file is required.'}), 400
    if not sig_b64:
        return jsonify({'error': 'Signature image is required.'}), 400
    input_path = os.path.join(UPLOAD_FOLDER, get_unique_filename('pdf'))
    file.save(input_path)
    try:
        import base64
        img_bytes = base64.b64decode(sig_b64)
        out_filename = f"signed_{uuid.uuid4().hex[:8]}.pdf"
        out_path = os.path.join(PROCESSED_FOLDER, out_filename)
        doc = fitz.open(input_path)
        page = doc[page_idx]
        rect = fitz.Rect(left, top, left + width, top + height)
        page.insert_image(rect, stream=img_bytes)
        doc.save(out_path)
        doc.close()
        return jsonify({'download_url': f'/download/{out_filename}'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try:
            os.remove(input_path)
        except:
            pass

@app.route('/ads/<path:filename>')
def serve_ad_snippets(filename):
    return send_from_directory(os.path.join(BASE_DIR, 'ads'), filename)

@app.route('/<path:slug>')
def dynamic_seo_page(slug):
    slug = slug.strip('/')
    
    # If the slug is actually a tool slug (like a legacy typo) that wasn't matched by specific routes
    if slug in SLUG_TO_FILE or slug in LEGACY_REDIRECTS:
        return serve_tool_page(slug)
        
    # Handle dynamic real working language/locale routes
    if is_valid_language_code(slug):
        index_path = os.path.join(BASE_DIR, 'index.html')
        with open(index_path, 'r', encoding='utf-8') as f:
            html = f.read()
        return apply_language_translations(html, slug)
        
    # Allow serving static files or explicit HTML from root
    root_file = os.path.join(BASE_DIR, slug)
    if os.path.exists(root_file) and os.path.isfile(root_file):
        return send_from_directory(BASE_DIR, slug)
        
    if not slug.endswith('.html'):
        html_file = os.path.join(BASE_DIR, slug + '.html')
        if os.path.exists(html_file) and os.path.isfile(html_file):
            return send_from_directory(BASE_DIR, slug + '.html')
    
    abort(404)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)



# Master Multi-Language Translation Map
DYNAMIC_LANG_MAP = {'es': {'title': 'i loves pdf | Herramientas PDF gratuitas en línea', 'h1': 'Todas las herramientas PDF que necesitas en un solo lugar', 'desc': 'Todas las herramientas PDF al alcance de tu mano. 100% GRATIS y fáciles de usar. Une, divide, comprime, convierte, rota, desbloquea y añade marcas de agua a tus archivos PDF con solo unos pocos clics.', 'all': 'Todos', 'org': 'Organizar PDF', 'opt': 'Optimizar PDF', 'conv': 'Convertir PDF', 'sec': 'Seguridad PDF', 'intel': 'Inteligencia PDF', 'merge': 'Unir PDF', 'split': 'Dividir PDF', 'compress': 'Comprimir PDF', 'w2p': 'Word a PDF', 'p2w': 'PDF a Word', 'p2pp': 'PDF a PowerPoint', 'p2ex': 'PDF a Excel', 'pp2p': 'PowerPoint a PDF', 'ex2p': 'Excel a PDF', 'j2p': 'JPG a PDF', 'p2j': 'PDF a JPG', 'rot': 'Rotar PDF', 'prot': 'Proteger PDF', 'unl': 'Desbloquear PDF', 'wm': 'Marca de agua', 'pn': 'Números de página', 'h2p': 'HTML a PDF', 'txt': 'Extraer texto', 'del': 'Eliminar páginas', 'pdfa': 'PDF a PDF/A', 'rep': 'Reparar PDF', 'ocr': 'OCR PDF', 'ai': 'Resumidor de IA', 'trans': 'Traducir PDF', 'sign': 'Firmar PDF', 'comp': 'Comparar PDF', 'login': 'Iniciar sesión', 'signup': 'Registrarse', 'home': 'Inicio', 'all_tools': 'Todas las herramientas PDF', 'convert_pdf': 'Convertir PDF'}, 'fr': {'title': 'i loves pdf | Outils PDF gratuits en ligne', 'h1': 'Tous les outils PDF dont vous avez besoin au même endroit', 'desc': 'Tous les outils PDF à portée de main. 100% GRATUITS et faciles à utiliser ! Fusionnez, divisez, compressez, convertissez, pivotez, déverrouillez et ajoutez un filigrane à vos PDFs en quelques clics.', 'all': 'Tous', 'org': 'Organiser PDF', 'opt': 'Optimiser PDF', 'conv': 'Convertir PDF', 'sec': 'Sécurité PDF', 'intel': 'Intelligence PDF', 'merge': 'Fusionner PDF', 'split': 'Diviser PDF', 'compress': 'Compresser PDF', 'w2p': 'Word en PDF', 'p2w': 'PDF en Word', 'p2pp': 'PDF en PowerPoint', 'p2ex': 'PDF en Excel', 'pp2p': 'PowerPoint en PDF', 'ex2p': 'Excel en PDF', 'j2p': 'JPG en PDF', 'p2j': 'PDF en JPG', 'rot': 'Pivoter PDF', 'prot': 'Protéger PDF', 'unl': 'Déverrouiller PDF', 'wm': 'Filigrane', 'pn': 'Numérotation', 'h2p': 'HTML en PDF', 'txt': 'Extraire le texte', 'del': 'Supprimer des pages', 'pdfa': 'PDF en PDF/A', 'rep': 'Réparer PDF', 'ocr': 'OCR PDF', 'ai': 'Résumeur IA', 'trans': 'Traduire PDF', 'sign': 'Signer PDF', 'comp': 'Comparer PDF', 'login': 'Connexion', 'signup': "S'inscrire", 'home': 'Accueil', 'all_tools': 'Tous les outils PDF', 'convert_pdf': 'Convertir PDF'}, 'de': {'title': 'i loves pdf | Kostenlose Online-PDF-Werkzeuge', 'h1': 'Jedes PDF-Werkzeug, das Sie benötigen, an einem Ort', 'desc': 'Alle PDF-Werkzeuge für Sie bereit. 100% KOSTENLOS und einfach zu bedienen! PDF-Dateien zusammenfügen, teilen, komprimieren, konvertieren, drehen, entsperren und Wasserzeichen hinzufügen.', 'all': 'Alle', 'org': 'PDF organisieren', 'opt': 'PDF optimieren', 'conv': 'PDF konvertieren', 'sec': 'PDF-Sicherheit', 'intel': 'PDF-Intelligenz', 'merge': 'PDF zusammenfügen', 'split': 'PDF teilen', 'compress': 'PDF komprimieren', 'w2p': 'Word in PDF', 'p2w': 'PDF in Word', 'p2pp': 'PDF in PowerPoint', 'p2ex': 'PDF in Excel', 'pp2p': 'PowerPoint in PDF', 'ex2p': 'Excel in PDF', 'j2p': 'JPG in PDF', 'p2j': 'PDF in JPG', 'rot': 'PDF drehen', 'prot': 'PDF schützen', 'unl': 'PDF entsperren', 'wm': 'Wasserzeichen', 'pn': 'Seitenzahlen', 'h2p': 'HTML in PDF', 'txt': 'Text extrahieren', 'del': 'Seiten löschen', 'pdfa': 'PDF in PDF/A', 'rep': 'PDF reparieren', 'ocr': 'OCR PDF', 'ai': 'KI-Zusammenfassung', 'trans': 'PDF übersetzen', 'sign': 'PDF unterschreiben', 'comp': 'PDF vergleichen', 'login': 'Einloggen', 'signup': 'Registrieren', 'home': 'Startseite', 'all_tools': 'Alle PDF-Werkzeuge', 'convert_pdf': 'PDF konvertieren'}, 'pt': {'title': 'i loves pdf | Ferramentas PDF online gratuitas', 'h1': 'Todas as ferramentas PDF necessárias em um só lugar', 'desc': "Todas as ferramentas PDF ao seu alcance. 100% GRATUITAS e fáceis de usar! Mesclar, dividir, comprimir, converter, rotacionar, desbloquear e adicionar marcas d'água aos seus PDFs.", 'all': 'Todos', 'org': 'Organizar PDF', 'opt': 'Otimizar PDF', 'conv': 'Converter PDF', 'sec': 'Segurança PDF', 'intel': 'Inteligência PDF', 'merge': 'Mesclar PDF', 'split': 'Dividir PDF', 'compress': 'Comprimir PDF', 'w2p': 'Word para PDF', 'p2w': 'PDF para Word', 'p2pp': 'PDF para PowerPoint', 'p2ex': 'PDF para Excel', 'pp2p': 'PowerPoint para PDF', 'ex2p': 'Excel para PDF', 'j2p': 'JPG para PDF', 'p2j': 'PDF para JPG', 'rot': 'Rotacionar PDF', 'prot': 'Proteger PDF', 'unl': 'Desbloquear PDF', 'wm': "Marca d'água", 'pn': 'Números de página', 'h2p': 'HTML para PDF', 'txt': 'Extrair texto', 'del': 'Excluir páginas', 'pdfa': 'PDF para PDF/A', 'rep': 'Reparar PDF', 'ocr': 'OCR PDF', 'ai': 'Resumidor de IA', 'trans': 'Traduzir PDF', 'sign': 'Assinar PDF', 'comp': 'Comparar PDF', 'login': 'Entrar', 'signup': 'Cadastrar', 'home': 'Início', 'all_tools': 'Todas as ferramentas PDF', 'convert_pdf': 'Converter PDF'}, 'hi': {'title': 'i loves pdf | मुफ़्त ऑनलाइन पीडीएफ उपकरण', 'h1': 'पीडीएफ पर काम करने के लिए सभी आवश्यक उपकरण एक जगह', 'desc': 'आपकी उंगलियों पर सभी पीडीएफ उपकरण। 100% मुफ़्त और उपयोग में आसान! कुछ ही क्लिक में पीडीएफ़ को मर्ज, स्प्लिट, कंप्रेस, कन्वर्ट, रोटेट, अनलॉक और वॉटरमार्क करें।', 'all': 'सभी', 'org': 'पीडीएफ व्यवस्थित करें', 'opt': 'पीडीएफ ऑप्टिमाइज़ करें', 'conv': 'पीडीएफ कन्वर्ट करें', 'sec': 'पीडीएफ सुरक्षा', 'intel': 'पीडीएफ इंटेलिजेंस', 'merge': 'पीडीएफ मर्ज करें', 'split': 'पीडीएफ विभाजित करें', 'compress': 'पीडीएफ कंप्रेस करें', 'w2p': 'वर्ड से पीडीएफ', 'p2w': 'पीडीएफ से वर्ड', 'p2pp': 'पीडीएफ से पावरपॉइंट', 'p2ex': 'पीडीएफ से एक्सेल', 'pp2p': 'पावरपॉइंट से पीडीएफ', 'ex2p': 'एक्सेल से पीडीएफ', 'j2p': 'जेपीजी से पीडीएफ', 'p2j': 'पीडीएफ से जेपीजी', 'rot': 'पीडीएफ घुमाएँ', 'prot': 'पीडीएफ सुरक्षित करें', 'unl': 'पीडीएफ अनलॉक करें', 'wm': 'वॉटरमार्क', 'pn': 'पेज नंबर', 'h2p': 'एचटीएमएल से पीडीएफ', 'txt': 'टेक्स्ट निकालें', 'del': 'पेजएं हटाएं', 'pdfa': 'पीडीएफ से पीडीएफ/ए', 'rep': 'पीडीएफ रिपेयर करें', 'ocr': 'ओसीआर पीडीएफ', 'ai': 'एआई सारांश', 'trans': 'अनुवाद करें', 'sign': 'हस्ताक्षर करें', 'comp': 'पीडीएफ तुलना करें', 'login': 'लॉगिन', 'signup': 'साइन अप', 'home': 'होम', 'all_tools': 'सभी पीडीएफ टूल्स', 'convert_pdf': 'कन्वर्ट पीडीएफ'}, 'ar': {'title': 'i loves pdf | أدوات PDF مجانية عبر الإنترنت', 'h1': 'جميع أدوات PDF التي تحتاجها في مكان واحد', 'desc': 'جميع أدوات PDF بين يديك. مجانية 100% وسهلة الاستخدام! دمج وتقسيم وضغط وتحويل وتدوير وحماية وحذف صفحات PDF بنقرات قليلة.', 'all': 'الكل', 'org': 'تنظيم PDF', 'opt': 'تحسين PDF', 'conv': 'تحويل PDF', 'sec': 'أمان PDF', 'intel': 'ذكاء PDF', 'merge': 'دمج PDF', 'split': 'تقسيم PDF', 'compress': 'ضغط PDF', 'w2p': 'Word إلى PDF', 'p2w': 'PDF إلى Word', 'p2pp': 'PDF إلى PowerPoint', 'p2ex': 'PDF إلى Excel', 'pp2p': 'PowerPoint إلى PDF', 'ex2p': 'Excel إلى PDF', 'j2p': 'JPG إلى PDF', 'p2j': 'PDF إلى JPG', 'rot': 'تدوير PDF', 'prot': 'حماية PDF', 'unl': 'فتح PDF', 'wm': 'علامة مائية', 'pn': 'أرقام الصفحات', 'h2p': 'HTML إلى PDF', 'txt': 'استخراج النص', 'del': 'حذف الصفحات', 'pdfa': 'PDF إلى PDF/A', 'rep': 'إصلاح PDF', 'ocr': 'التعرف الضوئي', 'ai': 'ملخص الذكاء الاصطناعي', 'trans': 'ترجمة PDF', 'sign': 'توقيع PDF', 'comp': 'مقارنة PDF', 'login': 'تسجيل الدخول', 'signup': 'إنشاء حساب', 'home': 'الرئيسية', 'all_tools': 'جميع أدوات PDF', 'convert_pdf': 'تحويل PDF'}, 'zh': {'title': 'i loves pdf | 免费在线 PDF 工具', 'h1': '您在一个地方所需的所有 PDF 工具', 'desc': '所有 PDF 工具触手可及。100% 免费且易于使用！只需点击几下即可合并、拆分、压缩、转换、旋转、解锁和加水印 PDF。', 'all': '全部', 'org': '组织 PDF', 'opt': '优化 PDF', 'conv': '转换 PDF', 'sec': 'PDF 安全', 'intel': 'PDF 智能', 'merge': '合并 PDF', 'split': '拆分 PDF', 'compress': '压缩 PDF', 'w2p': 'Word 转 PDF', 'p2w': 'PDF 转 Word', 'p2pp': 'PDF 转 PowerPoint', 'p2ex': 'PDF 转 Excel', 'pp2p': 'PowerPoint 转 PDF', 'ex2p': 'Excel 转 PDF', 'j2p': 'JPG 转 PDF', 'p2j': 'PDF 转 JPG', 'rot': '旋转 PDF', 'prot': '加密 PDF', 'unl': '解密 PDF', 'wm': '添加水印', 'pn': '添加页码', 'h2p': 'HTML 转 PDF', 'txt': '提取文本', 'del': '删除页面', 'pdfa': 'PDF 转 PDF/A', 'rep': '修复 PDF', 'ocr': 'OCR 识别', 'ai': 'AI 摘要', 'trans': '翻译 PDF', 'sign': '电子签名', 'comp': '对比 PDF', 'login': '登录', 'signup': '注册', 'home': '首页', 'all_tools': '所有 PDF 工具', 'convert_pdf': '转换 PDF'}, 'ja': {'title': 'i loves pdf | 無料のオンライン PDF ツール', 'h1': 'PDFの作業に必要なすべてのツールを1箇所に', 'desc': 'すべてのPDFツールが手元に。100%無料で使いやすい！数クリックでPDFの結合、分割、圧縮、変換、回転、保護解除が可能です。', 'all': 'すべて', 'org': 'PDF 整理', 'opt': 'PDF 最適化', 'conv': 'PDF 変換', 'sec': 'PDF セキュリティ', 'intel': 'PDF インテリジェンス', 'merge': 'PDF 結合', 'split': 'PDF 分割', 'compress': 'PDF 圧縮', 'w2p': 'Word から PDF', 'p2w': 'PDF から Word', 'p2pp': 'PDF から PowerPoint', 'p2ex': 'PDF から Excel', 'pp2p': 'PowerPoint から PDF', 'ex2p': 'Excel から PDF', 'j2p': 'JPG から PDF', 'p2j': 'PDF から JPG', 'rot': 'PDF 回転', 'prot': 'PDF 保護', 'unl': 'PDF ロック解除', 'wm': '透かし追加', 'pn': 'ページ番号', 'h2p': 'HTML から PDF', 'txt': 'テキスト抽出', 'del': 'ページ削除', 'pdfa': 'PDF から PDF/A', 'rep': 'PDF 修復', 'ocr': 'OCR 処理', 'ai': 'AI 要約', 'trans': 'PDF 翻訳', 'sign': 'PDF 署名', 'comp': 'PDF 比較', 'login': 'ログイン', 'signup': '新規登録', 'home': 'ホーム', 'all_tools': 'すべての PDF ツール', 'convert_pdf': 'PDF 変換'}, 'ru': {'title': 'i loves pdf | Бесплатные онлайн-инструменты PDF', 'h1': 'Все инструменты PDF, необходимые для работы, в одном месте', 'desc': 'Все инструменты PDF у вас под рукой. 100% БЕСПЛАТНО и просто в использовании! Объединяйте, разделяйте, сжимайте, конвертируйте, поворачивайте и защищайте PDF в несколько кликов.', 'all': 'Все', 'org': 'Упорядочить PDF', 'opt': 'Оптимизировать PDF', 'conv': 'Конвертировать PDF', 'sec': 'Безопасность PDF', 'intel': 'Интеллект PDF', 'merge': 'Объединить PDF', 'split': 'Разделить PDF', 'compress': 'Сжать PDF', 'w2p': 'Word в PDF', 'p2w': 'PDF в Word', 'p2pp': 'PDF в PowerPoint', 'p2ex': 'PDF в Excel', 'pp2p': 'PowerPoint в PDF', 'ex2p': 'Excel в PDF', 'j2p': 'JPG в PDF', 'p2j': 'PDF в JPG', 'rot': 'Повернуть PDF', 'prot': 'Защитить PDF', 'unl': 'Снять защиту PDF', 'wm': 'Водяной знак', 'pn': 'Номера страниц', 'h2p': 'HTML в PDF', 'txt': 'Извлечь текст', 'del': 'Удалить страницы', 'pdfa': 'PDF в PDF/A', 'rep': 'Восстановить PDF', 'ocr': 'Распознавание OCR', 'ai': 'ИИ Суммаризатор', 'trans': 'Перевести PDF', 'sign': 'Подписать PDF', 'comp': 'Сравнить PDF', 'login': 'Войти', 'signup': 'Регистрация', 'home': 'Главная', 'all_tools': 'Все инструменты PDF', 'convert_pdf': 'Конвертировать PDF'}, 'it': {'title': 'i loves pdf | Strumenti PDF online gratuiti', 'h1': 'Tutti gli strumenti PDF di cui hai bisogno in un unico posto', 'desc': 'Tutti gli strumenti PDF a portata di mano. 100% GRATUITI e facili da usare! Unisci, dividi, comprimi, converti, ruota e sblocca i tuoi PDF in pochi clic.', 'all': 'Tutti', 'org': 'Organizza PDF', 'opt': 'Ottimizza PDF', 'conv': 'Converti PDF', 'sec': 'Sicurezza PDF', 'intel': 'Intelligenza PDF', 'merge': 'Unisci PDF', 'split': 'Dividi PDF', 'compress': 'Comprimi PDF', 'w2p': 'Da Word a PDF', 'p2w': 'Da PDF a Word', 'p2pp': 'Da PDF a PowerPoint', 'p2ex': 'Da PDF ad Excel', 'pp2p': 'Da PowerPoint a PDF', 'ex2p': 'Da Excel a PDF', 'j2p': 'Da JPG a PDF', 'p2j': 'Da PDF a JPG', 'rot': 'Ruota PDF', 'prot': 'Proteggi PDF', 'unl': 'Sblocca PDF', 'wm': 'Filigrana', 'pn': 'Numeri di pagina', 'h2p': 'Da HTML a PDF', 'txt': 'Estrai testo', 'del': 'Elimina pagine', 'pdfa': 'Da PDF a PDF/A', 'rep': 'Ripara PDF', 'ocr': 'OCR PDF', 'ai': 'Riepilogo IA', 'trans': 'Traduci PDF', 'sign': 'Firma PDF', 'comp': 'Confronta PDF', 'login': 'Accedi', 'signup': 'Registrati', 'home': 'Home', 'all_tools': 'Tutti gli strumenti PDF', 'convert_pdf': 'Converti PDF'}}

INDIAN_LANGS = {
    'hi': {
        'title': 'i loves pdf | मुफ़्त ऑनलाइन पीडीएफ उपकरण (भारत)', 
        'h1': 'पीडीएफ पर काम करने के लिए सभी आवश्यक उपकरण एक जगह', 
        'desc': 'आपकी उंगलियों पर सभी पीडीएफ उपकरण। 100% मुफ़्त और उपयोग में आसान! भारत में कुछ ही क्लिक में पीडीएफ़ को मर्ज, स्प्लिट, कंप्रेस, कन्वर्ट, रोटेट, अनलॉक और वॉटरमार्क करें।', 
        'all': 'सभी', 'org': 'पीडीएफ व्यवस्थित करें', 'opt': 'पीडीएफ ऑप्टिमाइज़ करें', 'conv': 'पीडीएफ कन्वर्ट करें', 'sec': 'पीडीएफ सुरक्षा', 'intel': 'पीडीएफ इंटेलिजेंस', 
        'merge': 'पीडीएफ मर्ज करें', 'split': 'पीडीएफ विभाजित करें', 'compress': 'पीडीएफ कंप्रेस करें', 'w2p': 'वर्ड से पीडीएफ', 'p2w': 'पीडीएफ से वर्ड', 'p2pp': 'पीडीएफ से पावरपॉइंट', 
        'p2ex': 'पीडीएफ से एक्सेल', 'pp2p': 'पावरपॉइंट से पीडीएफ', 'ex2p': 'एक्सेल से पीडीएफ', 'j2p': 'जेपीजी से पीडीएफ', 'p2j': 'पीडीएफ से जेपीजी', 'rot': 'पीडीएफ घुमाएँ', 
        'prot': 'पीडीएफ सुरक्षित करें', 'unl': 'पीडीएफ अनलॉक करें', 'wm': 'वॉटरमार्क', 'pn': 'पेज नंबर', 'h2p': 'एचटीएमएल से पीडीएफ', 'txt': 'टेक्स्ट निकालें', 'del': 'पेज हटाएं', 
        'pdfa': 'पीडीएफ से पीडीएफ/ए', 'rep': 'पीडीएफ रिपेयर करें', 'ocr': 'ओसीआर पीडीएफ', 'ai': 'एआई सारांश', 'trans': 'अनुवाद करें', 'sign': 'हस्ताक्षर करें', 'comp': 'पीडीएफ तुलना करें', 
        'login': 'लॉगिन', 'signup': 'साइन अप', 'home': 'होम', 'all_tools': 'सभी पीडीएफ टूल्स', 'convert_pdf': 'कन्वर्ट पीडीएफ'
    },
    'bn': {
        'title': 'i loves pdf | বিনামূল্যে অনলাইন পিডিএফ টুলস (ভারত)', 
        'h1': 'আপনার প্রয়োজনীয় সমস্ত পিডিএফ টুলস এক জায়গায়', 
        'desc': 'আপনার নখদর্পণে সমস্ত পিডিএফ টুলস। ১০০% বিনামূল্যে এবং সহজে ব্যবহারযোগ্য! ভারতে মাত্র কয়েকটি ক্লিকে পিডিএফ মার্জ, স্প্লিট, কম্প্রেস, রূপান্তর, আনলক করুন।', 
        'all': 'সব', 'org': 'পিডিএফ সাজান', 'opt': 'পিডিএফ অপ্টিমাইজ করুন', 'conv': 'পিডিএফ রূপান্তর', 'sec': 'পিডিএফ নিরাপত্তা', 'intel': 'পিডিএফ বুদ্ধিমত্তা', 
        'merge': 'পিডিএফ মার্জ করুন', 'split': 'পিডিএফ স্প্লিট করুন', 'compress': 'পিডিএফ কম্প্রেস করুন', 'w2p': 'ওয়ার্ড থেকে পিডিএফ', 'p2w': 'পিডিএফ থেকে ওয়ার্ড', 'p2pp': 'পিডিএফ থেকে পাওয়ারপয়েন্ট', 
        'p2ex': 'পিডিএফ থেকে এক্সেল', 'pp2p': 'পাওয়ারপয়েন্ট থেকে পিডিএফ', 'ex2p': 'এক্সেল থেকে পিডিএফ', 'j2p': 'জেপিজি থেকে পিডিএফ', 'p2j': 'পিডিএফ থেকে জেপিজি', 'rot': 'পিডিএফ ঘোরান', 
        'prot': 'পিডিএফ সুরক্ষিত করুন', 'unl': 'পিডিএফ আনলক করুন', 'wm': 'ওয়াটারমার্ক', 'pn': 'পৃষ্ঠা নম্বর', 'h2p': 'এইচটিএমএল থেকে পিডিএফ', 'txt': 'টেক্সট বের করুন', 'del': 'পৃষ্ঠা মুছুন', 
        'pdfa': 'পিডিএফ থেকে পিডিএফ/এ', 'rep': 'পিডিএফ মেরামত করুন', 'ocr': 'ওসিআর পিডিএফ', 'ai': 'এআই সারাংশ', 'trans': 'পিডিএফ অনুবাদ করুন', 'sign': 'পিডিএফ স্বাক্ষর করুন', 'comp': 'পিডিএফ তুলনা করুন', 
        'login': 'লগইন', 'signup': 'নিবন্ধন করুন', 'home': 'হোম', 'all_tools': 'সমস্ত পিডিএফ টুলস', 'convert_pdf': 'পিডিএফ রূপান্তর করুন'
    },
    'ta': {
        'title': 'i loves pdf | இலவச ஆன்லைன் PDF கருவிகள் (இந்தியா)', 
        'h1': 'உங்களுக்குத் தேவையான அனைத்து PDF கருவிகளும் ஒரே இடத்தில்', 
        'desc': 'உங்கள் விரல் நுனியில் அனைத்து PDF கருவிகள். 100% இலவசம் மற்றும் பயன்படுத்த எளிதானது! இந்தியாவில் PDF களை இணைக்கவும், பிரிக்கவும், சுருக்கவும், மாற்றவும்.', 
        'all': 'அனைத்தும்', 'org': 'PDF ஒழுங்கமைக்கவும்', 'opt': 'PDF மேம்படுத்தவும்', 'conv': 'PDF மாற்றவும்', 'sec': 'PDF பாதுகாப்பு', 'intel': 'PDF நுண்ணறிவு', 
        'merge': 'PDF இணைக்கவும்', 'split': 'PDF பிரிக்கவும்', 'compress': 'PDF சுருக்கவும்', 'w2p': 'Word இல் இருந்து PDF', 'p2w': 'PDF இல் இருந்து Word', 'p2pp': 'PDF இல் இருந்து PowerPoint', 
        'p2ex': 'PDF இல் இருந்து Excel', 'pp2p': 'PowerPoint இல் இருந்து PDF', 'ex2p': 'Excel இல் இருந்து PDF', 'j2p': 'JPG இல் இருந்து PDF', 'p2j': 'PDF இல் இருந்து JPG', 'rot': 'PDF சுழற்றவும்', 
        'prot': 'PDF பாதுகாக்கவும்', 'unl': 'PDF திறக்கவும்', 'wm': 'வாட்டர்மார்க்', 'pn': 'பக்க எண்கள்', 'h2p': 'HTML இல் இருந்து PDF', 'txt': 'உரையைப் பிரித்தெடுக்கவும்', 'del': 'பக்கங்களை நீக்கவும்', 
        'pdfa': 'PDF இல் இருந்து PDF/A', 'rep': 'PDF சரிசெய்யவும்', 'ocr': 'OCR PDF', 'ai': 'AI சுருக்கம்', 'trans': 'PDF மொழிபெயர்க்கவும்', 'sign': 'PDF கையொப்பமிடவும்', 'comp': 'PDF ஒப்பிடவும்', 
        'login': 'உள்நுழைய', 'signup': 'பதிவு செய்க', 'home': 'முகப்பு', 'all_tools': 'அனைத்து PDF கருவிகள்', 'convert_pdf': 'PDF மாற்றவும்'
    },
    'te': {
        'title': 'i loves pdf | ఉచిత ఆన్‌లైన్ PDF సాధనాలు (భారతదేశం)', 
        'h1': 'మీకు కావలసిన అన్ని PDF సాధనాలు ఒకే చోట', 
        'desc': 'మీ చేతివేళ్ల వద్ద అన్ని PDF సాధనాలు. 100% ఉచితం మరియు ఉపయోగించడానికి సులభం! భారతదేశంలో PDF లను విలీనం చేయండి, విభజించండి, కుదించండి మరియు మార్చండి.', 
        'all': 'అన్నీ', 'org': 'PDF నిర్వహించండి', 'opt': 'PDF ఆప్టిమైజ్ చేయండి', 'conv': 'PDF మార్చండి', 'sec': 'PDF భద్రత', 'intel': 'PDF ఇంటెలిజెన్స్', 
        'merge': 'PDF విలీనం చేయండి', 'split': 'PDF విభజించండి', 'compress': 'PDF కుదించండి', 'w2p': 'Word నుండి PDF', 'p2w': 'PDF నుండి Word', 'p2pp': 'PDF నుండి PowerPoint', 
        'p2ex': 'PDF నుండి Excel', 'pp2p': 'PowerPoint నుండి PDF', 'ex2p': 'Excel నుండి PDF', 'j2p': 'JPG నుండి PDF', 'p2j': 'PDF నుండి JPG', 'rot': 'PDF తిప్పండి', 
        'prot': 'PDF రక్షించండి', 'unl': 'PDF అన్‌లాక్ చేయండి', 'wm': 'వాటర్‌మార్క్', 'pn': 'పేజీ సంఖ్యలు', 'h2p': 'HTML నుండి PDF', 'txt': 'వచనాన్ని సంగ్రహించండి', 'del': 'పేజీలను తొలగించండి', 
        'pdfa': 'PDF నుండి PDF/A', 'rep': 'PDF మరమ్మతు', 'ocr': 'OCR PDF', 'ai': 'AI సారాంశం', 'trans': 'PDF అనువదించండి', 'sign': 'PDF సంతకం చేయండి', 'comp': 'PDF సరిపోల్చండి', 
        'login': 'లాగిన్', 'signup': 'సైన్ అప్', 'home': 'హోమ్', 'all_tools': 'అన్ని PDF సాధనాలు', 'convert_pdf': 'PDF మార్చండి'
    },
    'mr': {
        'title': 'i loves pdf | मोफत ऑनलाइन पीडीएफ टूल्स (भारत)', 
        'h1': 'तुम्हाला आवश्यक असलेली सर्व पीडीएफ टूल्स एकाच ठिकाणी', 
        'desc': 'तुमच्या बोटांच्या टोकावर सर्व पीडीएफ टूल्स. 100% मोफत आणि वापरण्यास सोपे! भारतात पीडीएफ एकत्र करा, विभाजित करा, संकुचित करा आणि रूपांतरित करा.', 
        'all': 'सर्व', 'org': 'पीडीएफ आयोजित करा', 'opt': 'पीडीएफ ऑप्टिमाइझ करा', 'conv': 'पीडीएफ रूपांतरित करा', 'sec': 'पीडीएफ सुरक्षा', 'intel': 'पीडीएफ बुद्धिमत्ता', 
        'merge': 'पीडीएफ एकत्र करा', 'split': 'पीडीएफ विभाजित करा', 'compress': 'पीडीएफ संकुचित करा', 'w2p': 'वर्ड ते पीडीएफ', 'p2w': 'पीडीएफ ते वर्ड', 'p2pp': 'पीडीएफ ते पॉवरपॉइंट', 
        'p2ex': 'पीडीएफ ते एक्सेल', 'pp2p': 'पॉवरपॉइंट ते पीडीएफ', 'ex2p': 'एक्सेल ते पीडीएफ', 'j2p': 'जेपीजी ते पीडीएफ', 'p2j': 'पीडीएफ ते जेपीजी', 'rot': 'पीडीएफ फिरवा', 
        'prot': 'पीडीएफ सुरक्षित करा', 'unl': 'पीडीएफ अनलॉक करा', 'wm': 'वॉटरमार्क', 'pn': 'पृष्ठ क्रमांक', 'h2p': 'एचटीएमएल ते पीडीएफ', 'txt': 'मजकूर काढा', 'del': 'पृष्ठे हटवा', 
        'pdfa': 'पीडीएफ ते पीडीएफ/ए', 'rep': 'पीडीएफ दुरुस्त करा', 'ocr': 'ओसीआर पीडीएफ', 'ai': 'एआय सारांश', 'trans': 'पीडीएफ भाषांतरित करा', 'sign': 'पीडीएफ स्वाक्षरी करा', 'comp': 'पीडीएफ तुलना करा', 
        'login': 'लॉगिन', 'signup': 'साइन अप', 'home': 'होम', 'all_tools': 'सर्व पीडीएफ टूल्स', 'convert_pdf': 'पीडीएफ रूपांतरित करा'
    },
    'gu': {
        'title': 'i loves pdf | મફત ઓનલાઈન પીડીએફ ટૂલ્સ (ભારત)', 
        'h1': 'તમને જોઈતા તમામ પીડીએફ ટૂલ્સ એક જ જગ્યાએ', 
        'desc': 'તમારી આંગળીના વેઢે તમામ પીડીએફ ટૂલ્સ. 100% મફત અને ઉપયોગમાં સરળ! ભારતમાં પીડીએફ મર્જ કરો, વિભાજીત કરો, સંકુચિત કરો અને કન્વર્ટ કરો.', 
        'all': 'બધા', 'org': 'પીડીએફ ગોઠવો', 'opt': 'પીડીએફ ઓપ્ટિમાઇઝ કરો', 'conv': 'પીડીએફ કન્વર્ટ કરો', 'sec': 'પીડીએફ સુરક્ષા', 'intel': 'પીડીએફ બુદ્ધિ', 
        'merge': 'પીડીએફ મર્જ કરો', 'split': 'પીડીએફ વિભાજીત કરો', 'compress': 'પીડીએફ સંકુચિત કરો', 'w2p': 'વર્ડ થી પીડીએફ', 'p2w': 'પીડીએફ થી વર્ડ', 'p2pp': 'પીડીએફ થી પાવરપોઈન્ટ', 
        'p2ex': 'પીડીએફ થી એક્સેલ', 'pp2p': 'પાવરપોઈન્ટ થી પીડીએફ', 'ex2p': 'એક્સેલ થી પીડીએફ', 'j2p': 'જેપીજી થી પીડીએફ', 'p2j': 'પીડીએફ થી જેપીજી', 'rot': 'પીડીએફ ફેરવો', 
        'prot': 'પીડીએફ સુરક્ષિત કરો', 'unl': 'પીડીએફ અનલૉક કરો', 'wm': 'વોટરમાર્ક', 'pn': 'પૃષ્ઠ નંબરો', 'h2p': 'એચટીએમએલ થી પીડીએફ', 'txt': 'ટેક્સ્ટ કાઢો', 'del': 'પૃષ્ઠો કાઢી નાખો', 
        'pdfa': 'પીડીએફ થી પીડીએફ/એ', 'rep': 'પીડીએફ રિપેર કરો', 'ocr': 'ઓસીઆર પીડીએફ', 'ai': 'એઆઈ સારાંશ', 'trans': 'પીડીએફ અનુવાદ કરો', 'sign': 'પીડીએફ સહી કરો', 'comp': 'પીડીએફ સરખામણી કરો', 
        'login': 'લોગિન', 'signup': 'સાઇન અપ', 'home': 'હોમ', 'all_tools': 'તમામ પીડીએફ ટૂલ્સ', 'convert_pdf': 'પીડીએફ કન્વર્ટ કરો'
    },
    'kn': {
        'title': 'i loves pdf | ಉಚಿತ ಆನ್‌ಲೈನ್ PDF ಪರಿಕರಗಳು (ಭಾರತ)', 
        'h1': 'ನಿಮಗೆ ಅಗತ್ಯವಿರುವ ಎಲ್ಲಾ PDF ಪರಿಕರಗಳು ಒಂದೇ ಸ್ಥಳದಲ್ಲಿ', 
        'desc': 'ನಿಮ್ಮ ಬೆರಳ ತುದಿಯಲ್ಲಿ ಎಲ್ಲಾ PDF ಪರಿಕರಗಳು. 100% ಉಚಿತ ಮತ್ತು ಬಳಸಲು ಸುಲಭ! ಭಾರತದಲ್ಲಿ PDF ಗಳನ್ನು ವಿಲೀನಗೊಳಿಸಿ, ವಿಭಜಿಸಿ, ಕುಗ್ಗಿಸಿ ಮತ್ತು ಪರಿವರ್ತಿಸಿ.', 
        'all': 'ಎಲ್ಲಾ', 'org': 'PDF ಆಯೋಜಿಸಿ', 'opt': 'PDF ಆಪ್ಟಿಮೈಜ್ ಮಾಡಿ', 'conv': 'PDF ಪರಿವರ್ತಿಸಿ', 'sec': 'PDF ಭದ್ರತೆ', 'intel': 'PDF ಬುದ್ಧಿವಂತಿಕೆ', 
        'merge': 'PDF ವಿಲೀನಗೊಳಿಸಿ', 'split': 'PDF ವಿಭಜಿಸಿ', 'compress': 'PDF ಕುಗ್ಗಿಸಿ', 'w2p': 'Word ನಿಂದ PDF', 'p2w': 'PDF ನಿಂದ Word', 'p2pp': 'PDF ನಿಂದ PowerPoint', 
        'p2ex': 'PDF ನಿಂದ Excel', 'pp2p': 'PowerPoint ನಿಂದ PDF', 'ex2p': 'Excel ನಿಂದ PDF', 'j2p': 'JPG ನಿಂದ PDF', 'p2j': 'PDF ನಿಂದ JPG', 'rot': 'PDF ತಿರುಗಿಸಿ', 
        'prot': 'PDF ರಕ್ಷಿಸಿ', 'unl': 'PDF ಅನ್‌ಲಾಕ್ ಮಾಡಿ', 'wm': 'ವಾಟರ್‌ಮಾರ್ಕ್', 'pn': 'ಪುಟ ಸಂಖ್ಯೆಗಳು', 'h2p': 'HTML ನಿಂದ PDF', 'txt': 'ಪಠ್ಯವನ್ನು ಹೊರತೆಗೆಯಿರಿ', 'del': 'ಪುಟಗಳನ್ನು ಅಳಿಸಿ', 
        'pdfa': 'PDF ನಿಂದ PDF/A', 'rep': 'PDF ದುರಸ್ತಿ ಮಾಡಿ', 'ocr': 'OCR PDF', 'ai': 'AI ಸಾರಾಂಶ', 'trans': 'PDF ಭಾಷಾಂತರಿಸಿ', 'sign': 'PDF ಸಹಿ ಮಾಡಿ', 'comp': 'PDF ಹೋಲಿಕೆ ಮಾಡಿ', 
        'login': 'ಲಾಗಿನ್', 'signup': 'ಸೈನ್ ಅಪ್', 'home': 'ಮುಖಪುಟ', 'all_tools': 'ಎಲ್ಲಾ PDF ಪರಿಕರಗಳು', 'convert_pdf': 'PDF ಪರಿವರ್ತಿಸಿ'
    },
    'ur': {
        'title': 'i loves pdf | مفت آن لائن پی ڈی ایف ٹولز (بھارت)', 
        'h1': 'آپ کی ضرورت کے تمام پی ڈی ایف ٹولز ایک ہی جگہ پر', 
        'desc': 'آپ کی انگلیوں پر تمام پی ڈی ایف ٹولز۔ 100% مفت اور استعمال میں آسان! بھارت میں پی ڈی ایف کو ضم، تقسیم، کمپریس اور تبدیل کریں۔', 
        'all': 'تمام', 'org': 'پی ڈی ایف کو منظم کریں', 'opt': 'پی ڈی ایف کو بہتر بنائیں', 'conv': 'پی ڈی ایف کو تبدیل کریں', 'sec': 'پی ڈی ایف سیکیورٹی', 'intel': 'پی ڈی ایف انٹیلی جنس', 
        'merge': 'پی ڈی ایف ضم کریں', 'split': 'پی ڈی ایف تقسیم کریں', 'compress': 'پی ڈی ایف کمپریس کریں', 'w2p': 'ورڈ سے پی ڈی ایف', 'p2w': 'پی ڈی ایف سے ورڈ', 'p2pp': 'پی ڈی ایف سے پاورپوائنٹ', 
        'p2ex': 'پی ڈی ایف سے ایکسل', 'pp2p': 'پاورپوائنٹ سے پی ڈی ایف', 'ex2p': 'ایکسل سے پی ڈی ایف', 'j2p': 'جے پی جی سے پی ڈی ایف', 'p2j': 'پی ڈی ایف سے جے پی جی', 'rot': 'پی ڈی ایف گھمائیں', 
        'prot': 'پی ڈی ایف کو محفوظ بنائیں', 'unl': 'پی ڈی ایف ان لاک کریں', 'wm': 'واٹر مارک', 'pn': 'صفحہ نمبر', 'h2p': 'ایچ ٹی ایم ایل سے پی ڈی ایف', 'txt': 'متن نکالیں', 'del': 'صفحات حذف کریں', 
        'pdfa': 'پی ڈی ایف سے پی ڈی ایف/اے', 'rep': 'پی ڈی ایف کی مرمت', 'ocr': 'او سی آر پی ڈی ایف', 'ai': 'اے آئی خلاصہ', 'trans': 'پی ڈی ایف کا ترجمہ کریں', 'sign': 'پی ڈی ایف پر دستخط کریں', 'comp': 'پی ڈی ایف کا موازنہ کریں', 
        'login': 'لاگ ان', 'signup': 'سائن اپ', 'home': 'ہوم', 'all_tools': 'تمام پی ڈی ایف ٹولز', 'convert_pdf': 'پی ڈی ایف تبدیل کریں'
    },
    'pa': {
        'title': 'i loves pdf | ਮੁਫ਼ਤ ਆਨਲਾਈਨ PDF ਟੂਲ (ਭਾਰਤ)', 
        'h1': 'ਤੁਹਾਨੂੰ ਲੋੜੀਂਦੇ ਸਾਰੇ PDF ਟੂਲ ਇੱਕੋ ਥਾਂ \'ਤੇ', 
        'desc': 'ਤੁਹਾਡੀਆਂ ਉਂਗਲਾਂ \'ਤੇ ਸਾਰੇ PDF ਟੂਲ। 100% ਮੁਫ਼ਤ ਅਤੇ ਵਰਤਣ ਵਿੱਚ ਆਸਾਨ! ਭਾਰਤ ਵਿੱਚ ਕੁਝ ਹੀ ਕਲਿੱਕਾਂ ਨਾਲ PDF ਨੂੰ ਮਿਲਾਓ, ਵੰਡੋ, ਕੰਪ੍ਰੈਸ ਕਰੋ ਅਤੇ ਬਦਲੋ।', 
        'all': 'ਸਾਰੇ', 'org': 'PDF ਪ੍ਰਬੰਧਿਤ ਕਰੋ', 'opt': 'PDF ਅਨੁਕੂਲਿਤ ਕਰੋ', 'conv': 'PDF ਬਦਲੋ', 'sec': 'PDF ਸੁਰੱਖਿਆ', 'intel': 'PDF ਇੰਟੈਲੀਜੈਂਸ', 
        'merge': 'PDF ਮਿਲਾਓ', 'split': 'PDF ਵੰਡੋ', 'compress': 'PDF ਕੰਪ੍ਰੈਸ ਕਰੋ', 'w2p': 'Word ਤੋਂ PDF', 'p2w': 'PDF ਤੋਂ Word', 'p2pp': 'PDF ਤੋਂ PowerPoint', 
        'p2ex': 'PDF ਤੋਂ Excel', 'pp2p': 'PowerPoint ਤੋਂ PDF', 'ex2p': 'Excel ਤੋਂ PDF', 'j2p': 'JPG ਤੋਂ PDF', 'p2j': 'PDF ਤੋਂ JPG', 'rot': 'PDF ਘੁੰਮਾਓ', 
        'prot': 'PDF ਸੁਰੱਖਿਅਤ ਕਰੋ', 'unl': 'PDF ਅਨਲੌਕ ਕਰੋ', 'wm': 'ਵਾਟਰਮਾਰਕ', 'pn': 'ਪੰਨਾ ਨੰਬਰ', 'h2p': 'HTML ਤੋਂ PDF', 'txt': 'ਟੈਕਸਟ ਕੱਢੋ', 'del': 'ਪੰਨੇ ਮਿਟਾਓ', 
        'pdfa': 'PDF ਤੋਂ PDF/A', 'rep': 'PDF ਮੁਰੰਮਤ ਕਰੋ', 'ocr': 'OCR PDF', 'ai': 'AI ਸਾਰਾਂਸ਼', 'trans': 'PDF ਅਨੁਵਾਦ ਕਰੋ', 'sign': 'PDF ਦਸਤਖਤ ਕਰੋ', 'comp': 'PDF ਤੁਲਨਾ ਕਰੋ', 
        'login': 'ਲਾਗਿਨ', 'signup': 'ਸਾਈਨ ਅੱਪ', 'home': 'ਹੋਮ', 'all_tools': 'ਸਾਰੇ PDF ਟੂਲ', 'convert_pdf': 'PDF ਬਦਲੋ'
    },
    'ml': {
        'title': 'i loves pdf | സൗജന്യ ഓൺലൈൻ PDF ടൂളുകൾ (ഇന്ത്യ)', 
        'h1': 'നിങ്ങൾക്ക് ആവശ്യമായ എല്ലാ PDF ടൂളുകളും ഒരിടത്ത്', 
        'desc': 'നിങ്ങളുടെ വിരൽത്തുമ്പിൽ എല്ലാ PDF ടൂളുകളും. 100% സൗജന്യവും ഉപയോഗിക്കാൻ എളുപ്പവുമാണ്! ഇന്ത്യയിൽ PDF-കൾ ലയിപ്പിക്കുക, വിഭജിക്കുക, കംപ്രസ് ചെയ്യുക, പരിവർത്തനം ചെയ്യുക.', 
        'all': 'എല്ലാം', 'org': 'PDF സംഘടിപ്പിക്കുക', 'opt': 'PDF ഒപ്റ്റിമൈസ് ചെയ്യുക', 'conv': 'PDF പരിവർത്തനം ചെയ്യുക', 'sec': 'PDF സുരക്ഷ', 'intel': 'PDF ഇന്റലിജൻസ്', 
        'merge': 'PDF ലയിപ്പിക്കുക', 'split': 'PDF വിഭജിക്കുക', 'compress': 'PDF കംപ്രസ് ചെയ്യുക', 'w2p': 'Word-ൽ നിന്ന് PDF', 'p2w': 'PDF-ൽ നിന്ന് Word', 'p2pp': 'PDF-ൽ നിന്ന് PowerPoint', 
        'p2ex': 'PDF-ൽ നിന്ന് Excel', 'pp2p': 'PowerPoint-ൽ നിന്ന് PDF', 'ex2p': 'Excel-ൽ നിന്ന് PDF', 'j2p': 'JPG-ൽ നിന്ന് PDF', 'p2j': 'PDF-ൽ നിന്ന് JPG', 'rot': 'PDF തിരിക്കുക', 
        'prot': 'PDF സംരക്ഷിക്കുക', 'unl': 'PDF അൺലോക്ക് ചെയ്യുക', 'wm': 'വാട്ടർമാർക്ക്', 'pn': 'പേജ് നമ്പറുകൾ', 'h2p': 'HTML-ൽ നിന്ന് PDF', 'txt': 'വാചകം എക്‌സ്‌ട്രാക്‌റ്റ് ചെയ്യുക', 'del': 'പേജുകൾ ഇല്ലാതാക്കുക', 
        'pdfa': 'PDF-ൽ നിന്ന് PDF/A', 'rep': 'PDF നന്നാക്കുക', 'ocr': 'OCR PDF', 'ai': 'AI സംഗ്രഹം', 'trans': 'PDF വിവർത്തനം ചെയ്യുക', 'sign': 'PDF ഒപ്പിടുക', 'comp': 'PDF താരതമ്യം ചെയ്യുക', 
        'login': 'ലോഗിൻ', 'signup': 'സൈൻ അപ്പ്', 'home': 'ഹോം', 'all_tools': 'എല്ലാ PDF ടൂളുകളും', 'convert_pdf': 'PDF പരിവർത്തനം ചെയ്യുക'
    },
    'or': {
        'title': 'i loves pdf | ମାଗଣା ଅନଲାଇନ୍ ପିଡିଏଫ୍ ଟୁଲ୍ସ (ଭାରତ)', 
        'h1': 'ଆପଣଙ୍କର ସମସ୍ତ ପିଡିଏଫ୍ ଟୁଲ୍ ଗୋଟିଏ ସ୍ଥାନରେ', 
        'desc': 'ଆପଣଙ୍କ ଆଙ୍ଗୁଠି ଅଗରେ ସମସ୍ତ ପିଡିଏଫ୍ ଟୁଲ୍। ୧୦୦% ମାଗଣା ଏବଂ ବ୍ୟବହାର କରିବାକୁ ସହଜ! ଭାରତରେ ପିଡିଏଫ୍ ମିଶ୍ରଣ, ବିଭାଜନ, ସଙ୍କୋଚନ ଏବଂ ରୂପାନ୍ତର କରନ୍ତୁ।', 
        'all': 'ସମସ୍ତ', 'org': 'ପିଡିଏଫ୍ ସଜାନ୍ତୁ', 'opt': 'ପିଡିଏଫ୍ ଅପ୍ଟିମାଇଜ୍ କରନ୍ତୁ', 'conv': 'ପିଡିଏଫ୍ ରୂପାନ୍ତର କରନ୍ତୁ', 'sec': 'ପିଡିଏଫ୍ ସୁରକ୍ଷା', 'intel': 'ପିଡିଏଫ୍ ବୁଦ୍ଧିମତା', 
        'merge': 'ପିଡିଏଫ୍ ମିଶ୍ରଣ କରନ୍ତୁ', 'split': 'ପିଡିଏଫ୍ ବିଭାଜନ କରନ୍ତୁ', 'compress': 'ପିଡିଏଫ୍ ସଙ୍କୋଚନ କରନ୍ତୁ', 'w2p': 'ୱାର୍ଡ ରୁ ପିଡିଏଫ୍', 'p2w': 'ପିଡିଏଫ୍ ରୁ ୱାର୍ଡ', 'p2pp': 'ପିଡିଏଫ୍ ରୁ ପାୱାରପଏଣ୍ଟ', 
        'p2ex': 'ପିଡିଏଫ୍ ରୁ ଏକ୍ସେଲ', 'pp2p': 'ପାୱାରପଏଣ୍ଟ ରୁ ପିଡିଏଫ୍', 'ex2p': 'ଏକ୍ସେଲ ରୁ ପିଡିଏଫ୍', 'j2p': 'ଜେପିଜି ରୁ ପିଡିଏଫ୍', 'p2j': 'ପିଡିଏଫ୍ ରୁ ଜେପିଜି', 'rot': 'ପିଡିଏଫ୍ ଘୁରାନ୍ତୁ', 
        'prot': 'ପିଡିଏଫ୍ ସୁରକ୍ଷିତ କରନ୍ତୁ', 'unl': 'ପିଡିଏଫ୍ ଅନଲକ୍ କରନ୍ତୁ', 'wm': 'ୱାଟରମାର୍କ', 'pn': 'ପୃଷ୍ଠା ସଂଖ୍ୟା', 'h2p': 'ଏଚଟିଏମଏଲ ରୁ ପିଡିଏଫ୍', 'txt': 'ଟେକ୍ସଟ ବାହାର କରନ୍ତୁ', 'del': 'ପୃଷ୍ଠା ହଟାନ୍ତୁ', 
        'pdfa': 'ପିଡିଏଫ୍ ରୁ ପିଡିଏଫ୍/ଏ', 'rep': 'ପିଡିଏଫ୍ ମରାମତି କରନ୍ତୁ', 'ocr': 'ଓସିଆର ପିଡିଏଫ୍', 'ai': 'ଏଆଇ ସାରାଂଶ', 'trans': 'ପିଡିଏଫ୍ ଅନୁବାଦ କରନ୍ତୁ', 'sign': 'ପିଡିଏଫ୍ ସ୍ୱାକ୍ଷର କରନ୍ତୁ', 'comp': 'ପିଡିଏଫ୍ ତୁଳନା କରନ୍ତୁ', 
        'login': 'ଲଗଇନ୍', 'signup': 'ସାଇନ୍ ଅପ୍', 'home': 'ହୋମ୍', 'all_tools': 'ସମସ୍ତ ପିଡିଏଫ୍ ଟୁଲ୍ସ', 'convert_pdf': 'ପିଡିଏଫ୍ ରୂପାନ୍ତର କରନ୍ତୁ'
    }
}

DYNAMIC_LANG_MAP.update(INDIAN_LANGS)
from other_langs import OTHER_LANGS
DYNAMIC_LANG_MAP.update(OTHER_LANGS)

def apply_language_translations(html, lang):
    # Fix relative CSS link injected by previous script so it works on nested routes like /es/merge
    html = html.replace('href="./style.css', 'href="/style.css')
    
    if not lang or lang == 'en':
        return html
        
    # Get base 2-letter language code (e.g., 'es-mx' -> 'es')
    base_lang = lang.lower().split('-')[0]
    
    t_dict = DYNAMIC_LANG_MAP.get(base_lang)
    if not t_dict:
        return html
        
    # Localize internal links (href="/...")
    langs_prefixes = tuple(f"{l}/" for l in SUPPORTED_LANGS)
    static_prefixes = ('favicon', 'apple', 'site.webmanifest', 'logo', 'ad-', 'style.css', 'api/', 'download/') + langs_prefixes
    html = re.sub(
        r'href="/([^"]*)"', 
        lambda m: f'href="/{lang}/{m.group(1)}"' if not m.group(1).startswith(static_prefixes) else m.group(0), 
        html
    )
    
    # Fix the English language switcher link which gets incorrectly prefixed
    html = html.replace(f'<a href="/{lang}/">English</a>', '<a href="/">English</a>')
    
    # Localize inline JS navigation buttons (e.g., logo and Home buttons)
    html = re.sub(
        r'window\.location\.href=[\'"]\/[\'"]', 
        f'window.location.href="/{lang}/"', 
        html,
        flags=re.IGNORECASE
    )
        
    # Replace html lang tag
    html = re.sub(r'''<html\s+lang=["'][^"']*["']''', f'<html lang="{lang}"', html, flags=re.IGNORECASE)
    
    # Replace Titles & Meta Descriptions
    if 'title' in t_dict:
        html = re.sub(r'<title>.*?</title>', f'<title>{t_dict["title"]}</title>', html, flags=re.DOTALL)
    if 'desc' in t_dict:
        html = re.sub(r'''<meta\s+name=["']description["']\s+content=["'][^"']*["']''', f'<meta name="description" content="{t_dict["desc"]}"', html, flags=re.IGNORECASE)
        
    # Replace Hero H1 & Subtitle
    if 'h1' in t_dict:
        html = html.replace('Every tool you need to work with PDFs in one place', t_dict['h1'])
    if 'desc' in t_dict:
        html = html.replace('Every tool you need to use PDFs, at your fingertips. All are 100% FREE and easy to use! Merge, split, compress, convert, rotate, unlock and watermark PDFs with just a few clicks.', t_dict['desc'])
        
    # Replace Navigation & Category Tabs & Tools Grid
    replacements = [
        ('Merge PDF', t_dict.get('merge', 'Merge PDF')),
        ('Split PDF', t_dict.get('split', 'Split PDF')),
        ('Compress PDF', t_dict.get('compress', 'Compress PDF')),
        ('Word to PDF', t_dict.get('w2p', 'Word to PDF')),
        ('PDF to Word', t_dict.get('p2w', 'PDF to Word')),
        ('PDF to PowerPoint', t_dict.get('p2pp', 'PDF to PowerPoint')),
        ('PDF to Excel', t_dict.get('p2ex', 'PDF to Excel')),
        ('PowerPoint to PDF', t_dict.get('pp2p', 'PowerPoint to PDF')),
        ('Excel to PDF', t_dict.get('ex2p', 'Excel to PDF')),
        ('JPG to PDF', t_dict.get('j2p', 'JPG to PDF')),
        ('PDF to JPG', t_dict.get('p2j', 'PDF to JPG')),
        ('Rotate PDF', t_dict.get('rot', 'Rotate PDF')),
        ('Protect PDF', t_dict.get('prot', 'Protect PDF')),
        ('Unlock PDF', t_dict.get('unl', 'Unlock PDF')),
        ('Watermark', t_dict.get('wm', 'Watermark')),
        ('Page Numbers', t_dict.get('pn', 'Page Numbers')),
        ('Organize PDF', t_dict.get('org', 'Organize PDF')),
        ('HTML to PDF', t_dict.get('h2p', 'HTML to PDF')),
        ('Extract Text', t_dict.get('txt', 'Extract Text')),
        ('Delete Pages', t_dict.get('del', 'Delete Pages')),
        ('PDF to PDF/A', t_dict.get('pdfa', 'PDF to PDF/A')),
        ('Repair PDF', t_dict.get('rep', 'Repair PDF')),
        ('OCR PDF', t_dict.get('ocr', 'OCR PDF')),
        ('AI Summarizer', t_dict.get('ai', 'AI Summarizer')),
        ('Translate PDF', t_dict.get('trans', 'Translate PDF')),
        ('Sign PDF', t_dict.get('sign', 'Sign PDF')),
        ('Compare PDF', t_dict.get('comp', 'Compare PDF')),
        ('Log in', t_dict.get('login', 'Log in')),
        ('Sign up', t_dict.get('signup', 'Sign up')),
        ('Home', t_dict.get('home', 'Home')),
        ('All PDF Tools', t_dict.get('all_tools', 'All PDF Tools')),
        ('Convert PDF', t_dict.get('convert_pdf', 'Convert PDF'))
    ]
    
    for old_str, new_str in replacements:
        html = html.replace(old_str, new_str)
        
    # Advanced Regional SEO Schema Injection for Indian Languages
    indian_langs = ['hi', 'bn', 'ta', 'te', 'mr', 'gu', 'kn', 'ur', 'pa', 'ml', 'or']
    if base_lang in indian_langs:
        schema = f"""<script type="application/ld+json">
{{
  "@context": "https://schema.org",
  "@type": "WebSite",
  "name": "i loves pdfs",
  "url": "https://ilovespdfs.in/{lang}/",
  "inLanguage": "{lang}",
  "description": "{t_dict.get('desc', '')}"
}}
</script>
<script type="application/ld+json">
{{
  "@context": "https://schema.org",
  "@type": "Service",
  "name": "i loves pdfs",
  "areaServed": "IN",
  "description": "{t_dict.get('desc', '')}",
  "url": "https://ilovespdfs.in/{lang}/"
}}
</script>"""
        html = re.sub(r'</head>', schema + '\n</head>', html, flags=re.IGNORECASE)
        
    return html
